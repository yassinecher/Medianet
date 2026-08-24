# -*- coding: utf-8 -*-
"""Build the French, Scrum-restructured (release-grouped) Medianet Incubateur PFE report."""
import os, copy
from PIL import Image
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SOURCE_DOCX = os.path.join(HERE, "RAPPORT_MEDIANET_INCUBATEUR.docx")   # original report (in-body figures)
COVER_DOCX  = os.path.join(HERE, "Page Garde Rapport Stage(Fr).docx")   # official ESPRIT page de garde + validation form
STYLE_DOCX  = os.path.join(HERE, "fichiertye2023.docx")                 # ESPRIT style template (fonts + title styles)
MEDIA = os.path.join(HERE, "_scrum_media")
OUT = os.environ.get("SCRUM_OUT", os.path.join(HERE, "RAPPORT_MEDIANET_INCUBATEUR_SCRUM_FR.docx"))

# Extract in-body figure images from the original .docx once (self-contained, re-runnable).
if not os.path.isdir(MEDIA) or not os.path.exists(os.path.join(MEDIA, "image1.png")):
    import zipfile
    os.makedirs(MEDIA, exist_ok=True)
    with zipfile.ZipFile(SOURCE_DOCX) as z:
        for n in z.namelist():
            if n.startswith("word/media/"):
                b = os.path.basename(n)
                if b:
                    with z.open(n) as fsrc, open(os.path.join(MEDIA, b), "wb") as fdst:
                        fdst.write(fsrc.read())

HEADSH = "D9D9D9"
CONTENT_W = 16.0  # cm

# Generate the Scrum release-board planning figure if missing (replaces the old Gantt).
if not os.path.exists(os.path.join(HERE, "planning_scrum.png")):
    import runpy
    runpy.run_path(os.path.join(HERE, "build_planning_scrum.py"))

# ---- Cover title/author values that fill the ESPRIT page de garde placeholders ----
TITLE_LINE = "Conception et développement d'une plateforme de gestion de programmes d'incubation : Medianet Incubateur"
COVER_FILL = {
    "2022 - 2023": "2025 - 2026",
    "INFORMATIQUE": "ARCHITECTURE LOGICIELLE",
    "SUJET DU RAPPORT DE STAGE": TITLE_LINE,
    "LOGO ENTREPRISE": "MEDIANET",
    "Réalisé par: Prénom Nom": "Réalisé par : Yassine Cherni",
    "Réalisé par : Prénom Nom": "Réalisé par : Yassine Cherni",
    "Encadré par:": "Encadré par :",
    "Encadrant ESPRIT:": "Encadrant ESPRIT : Slim Guermazi",
    "Encadrant Entreprise:": "Encadrant Entreprise : Yassine Mechich",
}

def _norm(s):
    return " ".join((s or "").split()).replace(" :", ":")

# =========================================================================
# BASE = official ESPRIT page de garde; import the style template's styles.
# =========================================================================
doc = Document(COVER_DOCX)

# import all named styles (fonts + title styles) from the ESPRIT style template
_bs = doc.styles.element
_ts = Document(STYLE_DOCX).styles.element
for _c in list(_bs):
    _bs.remove(_c)
for _c in list(_ts):
    _bs.append(copy.deepcopy(_c))

# Give the unnumbered chapter/title style ('titre plus') a top outline level so the
# Table of Contents ( \o "1-3" ) captures chapters alongside Heading 2/3 sections.
_tp = doc.styles['titre plus'].element
_tp_pPr = _tp.find(qn('w:pPr'))
if _tp_pPr is None:
    _tp_pPr = OxmlElement('w:pPr'); _tp.insert(0, _tp_pPr)
for _old in _tp_pPr.findall(qn('w:outlineLvl')):
    _tp_pPr.remove(_old)
_ol = OxmlElement('w:outlineLvl'); _ol.set(qn('w:val'), '0'); _tp_pPr.append(_ol)

# fill the page-de-garde placeholders (both DrawingML text boxes and VML fallbacks)
body = doc.element.body
_children = list(body)
_cover_el = _children[40]          # the paragraph carrying the overlaid title/author text boxes
for _p in _cover_el.findall('.//' + qn('w:p')):
    _txts = _p.findall('.//' + qn('w:t'))
    if not _txts:
        continue
    _full = _norm("".join(t.text or "" for t in _txts))
    _repl = None
    for k, v in COVER_FILL.items():
        if _norm(k) == _full:
            _repl = v
            break
    if _repl is not None:
        _txts[0].text = _repl
        for _t in _txts[1:]:
            _t.text = ""

# reorder: cover page first (elems 9..40), then the validation form (elems 0..8)
_validation = _children[0:9]
_cover      = _children[9:41]
_sectpr     = _children[41]
for _c in _children:
    body.remove(_c)
for _c in _cover:
    body.append(_c)
_pb = OxmlElement('w:p'); _r = OxmlElement('w:r'); _br = OxmlElement('w:br'); _br.set(qn('w:type'), 'page')
_r.append(_br); _pb.append(_r); body.append(_pb)
for _c in _validation:
    body.append(_c)
body.append(_sectpr)   # keep sectPr last; add_* inserts new content before it

# The cover/validation paragraphs were laid out for the garde doc's own (tight) Normal.
# The imported template Normal has 1.5 line + 6pt spacing + 1cm indent, which would inflate
# the ~30 cover spacer paragraphs onto a 2nd page. Reset the pre-existing paragraphs to
# compact spacing so the official page de garde stays on a single page.
def _compact(p_el):
    pPr = p_el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); p_el.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing'); pPr.append(sp)
    sp.set(qn('w:before'), '0'); sp.set(qn('w:after'), '0')
    sp.set(qn('w:line'), '240'); sp.set(qn('w:lineRule'), 'auto')
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        ind.set(qn('w:firstLine'), '0'); ind.set(qn('w:left'), '0')
for _p in body.findall(qn('w:p')):
    _compact(_p)

# A4 + 2.5 cm margins on the single section
sec = doc.sections[-1]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
    setattr(sec, m, Cm(2.5))

# Caption: small italic centered (respect template font otherwise)
_cap = doc.styles['Caption']
_cap.font.size = Pt(10); _cap.font.italic = True
_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
_cap.paragraph_format.first_line_indent = Cm(0)
_cap.paragraph_format.space_after = Pt(10)

# ---------- helpers ----------
def _set_updatefields():
    settings = doc.settings.element
    el = OxmlElement('w:updateFields'); el.set(qn('w:val'), 'true')
    settings.append(el)

def rich(p, text):
    """Parse **bold**, *italic*, `code` into runs."""
    import re
    i = 0
    tokens = re.split(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            r = p.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = p.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(10.5)
        else:
            p.add_run(tok)

def P(text, align='justify', indent=True, space_after=None, size=None, bold=False, center=False):
    # Inherits the template's configured Normal (Times New Roman 12, 1.5 line, 1 cm first-line
    # indent, justified). Only override when explicitly asked.
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if center or align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if not indent:
            pf.first_line_indent = Cm(0)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    rich(p, text)
    if bold:
        for r in p.runs:
            r.bold = True
    if size:
        for r in p.runs:
            r.font.size = Pt(size)
    return p

def _suppress_num(p):
    """Cancel the template's automatic heading numbering (text carries its own numbers)."""
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl'); ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId'); numId.set(qn('w:val'), '0')
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)

def H1(text):
    # 'titre plus' = the ESPRIT template's unnumbered top-level title style
    doc.add_paragraph(text, style='titre plus')

def H2(text):
    p = doc.add_paragraph(text, style='Heading 2'); _suppress_num(p)

def H3(text):
    p = doc.add_paragraph(text, style='Heading 3'); _suppress_num(p)

def bullet(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Cm(0.9); pf.first_line_indent = Cm(-0.4)
    pf.space_after = Pt(3)
    p.add_run("•\t")
    rich(p, text)

def numbered(text):
    bullet(text)

def pagebreak():
    p = doc.add_paragraph()
    r = p.add_run(); r.add_break(WD_BREAK.PAGE)

def _seq(p, kind):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), f' SEQ {kind} \\* ARABIC ')
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = '1'; r.append(t); fld.append(r)
    p._p.append(fld)

def caption(kind, text):
    p = doc.add_paragraph(style='Caption')
    p.add_run(f'{kind} ')
    _seq(p, kind)
    run = p.add_run(f' \u2014 {text}')

def figure(imgfile, cap_text, max_w=CONTENT_W, max_h=20.0):
    path = imgfile if os.path.isabs(imgfile) else os.path.join(MEDIA, imgfile)
    w, h = Image.open(path).size
    ar = h / w
    width_cm = max_w
    if width_cm * ar > max_h:
        width_cm = max_h / ar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    caption('Figure', cap_text)

def _shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def table(rows, widths=None, header=True, cap_text=None, font_size=10.5):
    ncols = len(rows[0])
    t = doc.add_table(rows=0, cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths is None:
        widths = [CONTENT_W / ncols] * ncols
    # scale widths to content
    scale = CONTENT_W / sum(widths)
    widths = [w * scale for w in widths]
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cell = cells[ci]
            cell.width = Cm(widths[ci])
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.first_line_indent = Cm(0)
            rich(para, str(val))
            for r in para.runs:
                r.font.size = Pt(font_size)
                if ri == 0 and header:
                    r.bold = True
            if ri == 0 and header:
                _shade(cell, HEADSH)
    # enforce column widths on cells
    for row in t.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = Cm(widths[ci])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    if cap_text:
        caption('Tableau', cap_text)
    return t

def toc(kind='TOC'):
    """kind: 'TOC' (contents), 'FIG', 'TAB'"""
    instr = {
        'TOC': r'TOC \o "1-3" \h \z \u',
        'FIG': r'TOC \h \z \c "Figure"',
        'TAB': r'TOC \h \z \c "Tableau"',
    }[kind]
    p = doc.add_paragraph()
    r = p.add_run()
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = ' ' + instr + ' '
    fldSep = OxmlElement('w:fldChar'); fldSep.set(qn('w:fldCharType'), 'separate')
    tr = OxmlElement('w:r'); tt = OxmlElement('w:t'); tt.text = 'Mettre à jour ce champ (F9) dans Word.'; tr.append(tt)
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    r._r.append(fldBegin); r._r.append(instrText); r._r.append(fldSep)
    p._p.append(tr); r2 = p.add_run(); r2._r.append(fldEnd)

def fullpage_image(imgfile):
    path = os.path.join(MEDIA, imgfile)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(path, height=Cm(25.5))

# =========================================================================
# FRONT MATTER
# (the ESPRIT page de garde + the validation-du-dépôt form come from the base
#  document above; content below is appended after them, before the sectPr)
# =========================================================================
pagebreak()

# --- Dédicace ---
H1("Dédicace")
for line in [
    "À mes parents, pour leur soutien inconditionnel et tout ce qu'ils ont sacrifié en chemin,",
    "À mes enseignants, qui m'ont transmis leur savoir avec patience et rigueur,",
    "À mes amis et collègues, pour leurs encouragements et leur présence,",
    "À toutes celles et ceux qui, directement ou indirectement, ont contribué à ce travail,",
    "Je dédie ce modeste accomplissement.",
]:
    P(line, center=True, indent=False, space_after=8)
P("— Yassine Cherni", center=True, indent=False, bold=True)
pagebreak()

# --- Remerciements ---
H1("Remerciements")
P("Je tiens tout d'abord à remercier **Medianet** de m'avoir accueilli et de m'avoir offert "
  "l'opportunité de réaliser ce projet de fin d'études dans un environnement professionnel "
  "réellement stimulant.")
P("Mes sincères remerciements vont à mon encadrant entreprise, **Yassine Mechich**, pour sa "
  "disponibilité, ses conseils avisés et le suivi rigoureux qu'il m'a accordé tout au long de ce stage.")
P("Je remercie également mon encadrant académique, **Slim Guermazi**, pour son accompagnement "
  "méthodologique, ses retours constructifs et sa disponibilité malgré un emploi du temps chargé.")
P("J'adresse mes remerciements aux membres du jury pour l'honneur qu'ils me font en acceptant "
  "d'évaluer ce travail.")
P("Enfin, je remercie l'ensemble du corps enseignant de l'**ESPRIT** pour la qualité de la "
  "formation reçue durant mes études, ainsi que ma famille et mes amis pour leur soutien constant.")
pagebreak()

# --- Tables ---
H1("Table des matières")
toc('TOC')
pagebreak()
H1("Table des figures")
toc('FIG')
pagebreak()
H1("Liste des tableaux")
toc('TAB')
pagebreak()

# =========================================================================
# INTRODUCTION GÉNÉRALE
# =========================================================================
H1("Introduction générale")
P("L'économie numérique a fait de la **startup** un moteur central de création de valeur et "
  "d'emploi, et des **incubateurs** les institutions chargées d'accompagner les projets, de "
  "l'idée jusqu'au marché. En Tunisie, l'adoption du *Startup Act* et la multiplication des "
  "structures d'incubation et d'accélération ont fait croître, année après année, le nombre de "
  "programmes à gérer. Chaque programme met en scène une chorégraphie relativement complexe "
  "d'acteurs et d'événements : un appel à candidatures, la réception et le tri des dossiers, la "
  "constitution des jurys, les sessions de présélection et de pitch, l'intégration des startups "
  "sélectionnées, plusieurs mois de mentorat et de formation, et enfin une journée de "
  "démonstration devant des investisseurs.")
P("Il y a pourtant là un paradoxe : ces mêmes incubateurs, qui accompagnent les startups dans "
  "leur transformation numérique, gèrent le plus souvent leurs propres programmes de façon "
  "assez **artisanale** — formulaires de candidature éparpillés dans des outils tiers, grilles "
  "de sélection tenues dans des tableurs, jurys contactés par courriel, agendas maintenus à la "
  "main, communications perdues quelque part dans une boîte de réception. Résultat : le "
  "processus ralentit, les erreurs s'accumulent, la consolidation des résultats devient une "
  "corvée et la traçabilité fait défaut — précisément au moment où les financeurs et les "
  "partenaires réclament de la transparence et des résultats mesurables.")
P("C'est dans ce contexte que ce projet de fin d'études a été réalisé, au sein de **Medianet** : "
  "**concevoir et développer une plateforme intégrée de gestion des programmes d'incubation**, "
  "baptisée **Medianet Incubateur**. L'idée est d'offrir aux responsables de programme un "
  "environnement unique, cohérent et sécurisé pour concevoir un programme, recevoir et évaluer "
  "les candidatures, dérouler les sessions de son *parcours*, et gérer — tout en les informant — "
  "l'ensemble des personnes impliquées. Porteurs de projet, jurys et grand public bénéficient, "
  "quant à eux, d'une expérience d'accès adaptée à leur rôle.")
P("Ce projet ne se limite pas à numériser des pratiques existantes ; il introduit quelques "
  "idées qui façonnent réellement la solution. D'abord, un **modèle de session unifié et "
  "extensible**, capable de représenter n'importe quelle étape d'un programme — candidature, "
  "présélection, journée de pitch, intégration, incubation, démo day, formation — sans jamais "
  "toucher au schéma de base de données. Ensuite, une dimension de **visibilité** qui distingue "
  "sessions publiques, internes et privées. Une **couche de validation centralisée**, aussi, qui "
  "garantit la cohérence des dates et l'intégrité du parcours. Et enfin, une **évaluation "
  "assistée par IA** qui aide les jurys à noter les candidatures plus objectivement, au regard "
  "des critères propres à chaque programme.")
P("Le problème central que ce projet cherche à résoudre peut se formuler en une seule question : "
  "*comment concevoir une plateforme unique capable de modéliser un programme d'incubation comme "
  "un parcours structuré et visuel de sessions typées et contrôlées en accès, tout en "
  "garantissant la fiabilité des données, la sécurité d'accès fondée sur les rôles et la "
  "traçabilité des communications ?*")
P("Une seconde question, tout aussi centrale, a émergé à mesure que le projet avançait : jusqu'à "
  "quel point l'évaluation d'un pitch peut-elle être rendue objective, reproductible et utile à "
  "la personne évaluée — sans que la vidéo ne quitte jamais l'infrastructure de l'incubateur ? "
  "Demander à un modèle de langage de « noter ce pitch » est trivial et, comme ce rapport le "
  "montrera, peu fiable : il produit une prose assurée sans rien derrière. L'approche défendue "
  "ici est différente. Tout ce qui peut être mesuré l'est à partir du signal lui-même — débit de "
  "parole, intensité sonore et sa dynamique, vrais silences issus de la forme d'onde, "
  "hésitations alignées mot à mot — et le modèle de langage n'est ensuite sollicité que pour "
  "juger au regard de ces mesures, selon une grille calibrée sur le format du pitch. Ce qui ne "
  "peut être mesuré de façon fiable, la présence physique en particulier, est explicitement "
  "rétrogradé au rang d'indice faible plutôt que présenté comme un fait. L'ensemble du pipeline "
  "est auto-hébergé : transcription, analyse audio et vision s'exécutent localement, de sorte "
  "que le pitch confidentiel d'un porteur ne quitte jamais les lieux.")
P("Le corps de ce rapport suit une organisation **agile, orientée Scrum**. Le **chapitre 1** "
  "pose le cadre général — organisme d'accueil, existant, solution proposée et méthodologie "
  "Scrum retenue. Le **chapitre 2** établit les fondements théoriques sur lesquels repose "
  "l'analyse de pitch. Le **chapitre 3** correspond au *Sprint 0* : spécification des besoins, "
  "acteurs, backlog produit, architecture globale et environnement de travail. Les **chapitres "
  "4 à 6** déroulent les trois *releases* successives, chacune présentant son backlog, son "
  "analyse, sa conception, sa réalisation et sa revue de sprint. Le **chapitre 7** rassemble la "
  "validation, les tests et le déploiement. Une conclusion générale clôt le rapport et ouvre sur "
  "les perspectives.")
pagebreak()

# =========================================================================
# CHAPITRE 1
# =========================================================================
H1("Chapitre 1 — Cadre général du projet")
H2("Introduction")
P("Avant d'entrer dans la plateforme elle-même, il convient de planter le décor. Ce chapitre "
  "présente l'organisme d'accueil, le problème que le stage devait traiter et un premier aperçu "
  "de ce qui a été construit — suivi d'un examen de ce qui existe déjà, en Tunisie et ailleurs, "
  "et de ses limites. Il se termine par la méthodologie de développement retenue, **Scrum**, et "
  "par les raisons de ce choix.")

H2("1.1 Organisme d'accueil")
P("Cette section présente l'organisme qui a accueilli le projet : sa présentation générale, ses "
  "domaines d'activité, son organisation interne et le département au sein duquel le stage s'est "
  "déroulé.")
H3("1.1.1 Présentation générale")
P("**Medianet** est une entreprise tunisienne spécialisée dans les services numériques, le web "
  "et les solutions logicielles. Depuis sa création en 1998, l'entreprise accompagne des "
  "organisations publiques et privées dans leur transformation numérique, notamment à travers le "
  "développement web, l'hébergement, le marketing digital, l'e-business et l'ingénierie logicielle.")
figure('image2.png', "Logo de l'organisme d'accueil")
P("Le tableau ci-dessous résume les informations clés de l'organisme d'accueil.")
table([
    ["Rubrique", "Valeur"],
    ["Raison sociale", "Medianet"],
    ["Forme juridique", "SARL"],
    ["Année de création", "1998"],
    ["Siège social", "Ariana, Tunisie"],
    ["Secteur", "Services numériques / Web / Logiciel"],
    ["Activités principales", "Web, hébergement, marketing digital, ingénierie logicielle"],
    ["Effectif", "50 – 200 employés"],
    ["Site web", "www.medianet.tn"],
], widths=[5, 11], cap_text="Fiche signalétique de l'organisme d'accueil")

H3("1.1.2 Domaines d'activité")
P("Concrètement, l'organisme d'accueil est actif dans plusieurs domaines :")
bullet("le **développement web** et les applications sur mesure pour des clients publics et privés ;")
bullet("l'**hébergement** et les services managés pour les solutions numériques ;")
bullet("le **marketing digital** et l'accompagnement de la présence en ligne ;")
bullet("le **soutien à l'innovation**, en particulier l'aide à la structuration de programmes "
       "d'incubation et d'accélération — précisément le cadre dans lequel s'inscrit ce projet de "
       "fin d'études.")

H3("1.1.3 Organisation de l'entreprise")
P("[Insérer ici l'organigramme de l'entreprise]", center=True, indent=False)
caption('Figure', "Organigramme de l'organisme d'accueil")
P("L'organigramme ci-dessus présente la structure de l'entreprise. Elle est organisée en "
  "plusieurs départements, dont l'un a accueilli ce projet de fin d'études.")

H3("1.1.4 Département d'accueil")
P("Le stage s'est déroulé au sein du département **ingénierie logicielle / innovation**, sous "
  "l'encadrement de **Yassine Mechich**. C'est l'équipe qui conçoit et développe les applications "
  "web de l'entreprise, aussi bien pour un usage interne que pour ses clients.")
P("Le tableau ci-dessous résume le cadre administratif du stage.")
table([
    ["Rubrique", "Valeur"],
    ["Nature", "Projet de fin d'études"],
    ["Durée", "6 mois"],
    ["Période", "01/02/2026 – 31/07/2026"],
    ["Encadrant académique", "Slim Guermazi"],
    ["Encadrant entreprise", "Yassine Mechich"],
    ["Spécialité", "Ingénierie de l'architecture logicielle"],
], widths=[5, 11])

H2("1.2 Présentation du projet")
P("Ce projet consiste à concevoir et développer **Medianet Incubateur**, une plateforme web "
  "dont l'objectif est de numériser l'intégralité du cycle de vie d'un programme d'incubation de "
  "startups. Elle réunit, dans un outil unique adapté à chaque rôle, la conception visuelle d'un "
  "programme sous forme de parcours de sessions typées et contrôlées en visibilité, les "
  "candidatures en ligne, l'évaluation par le jury assistée par intelligence artificielle, la "
  "gestion des participants et des organisations, et des communications réellement traçables.")
figure('image3.png', "Vue d'ensemble de la solution proposée")
P("Comme le montre la figure ci-dessus, la solution repose sur deux interfaces distinctes : un "
  "**back-office**, destiné aux administrateurs et aux responsables de programme, et un "
  "**front-office** à destination des porteurs de projet, des jurys, des membres d'organisations "
  "et, plus largement, du grand public.")

H2("1.3 Étude de l'existant")
P("Concevoir une solution pertinente suppose d'abord de comprendre les pratiques en place et "
  "leurs limites. Cette section décrit d'abord la manière dont les incubateurs gèrent aujourd'hui "
  "leurs programmes, puis en dégage une critique qui motive les choix du projet.")
H3("1.3.1 Description de l'existant")
P("Que ce soit chez l'organisme d'accueil ou chez des acteurs comparables de l'écosystème "
  "tunisien, la gestion des programmes d'incubation repose encore largement sur un assemblage "
  "d'outils génériques :")
bullet("des formulaires de candidature construits avec des outils tiers, tels que Typeform ou Google Forms ;")
bullet("des grilles de sélection et de notation qui vivent dans des tableurs (Excel, Google Sheets) ;")
bullet("des agendas tenus à la main dans des calendriers partagés ;")
bullet("des communications gérées par courriel et messagerie instantanée, sans véritable archive centralisée ;")
bullet("des rapports d'avancement qu'il faut recalculer manuellement à chaque revue.")
P("C'est le cas de plusieurs incubateurs tunisiens — **Flat6Labs Tunis**, **Wiki Start Up**, "
  "**B@Labs (BIAT)**, **Impact Partner**, **The Dot**, **Redstart Tunisie** — qui publient leurs "
  "appels à candidatures et animent leurs cohortes de manière assez similaire : une page de "
  "présentation, un formulaire externe et un suivi qui reste manuel en interne. Le portail "
  "national **Startup Tunisia / Startup Act** suit une autre logique, puisqu'il se concentre sur "
  "la labellisation juridique des startups plutôt que sur la gestion opérationnelle du parcours "
  "d'un programme.")
P("Le tableau ci-dessous compare, fonctionnalité par fonctionnalité, ce que couvrent "
  "généralement ces solutions face aux ambitions de ce projet.")
table([
    ["Fonctionnalité", "Portails d'incubateurs génériques", "Portail national de labellisation", "Medianet Incubateur"],
    ["Présentation publique de programmes", "Oui", "Partielle", "Oui"],
    ["Formulaire de candidature en ligne", "Oui (souvent externe)", "Non", "Oui (par programme)"],
    ["Parcours visuel du programme", "Non", "Non", "Oui"],
    ["Sessions typées avec activités", "Non", "Non", "Oui"],
    ["Contrôle de visibilité des sessions", "Non", "Non", "Oui"],
    ["Évaluation par jury (critères pondérés)", "Rare / externe", "Non", "Oui"],
    ["Notation assistée par IA", "Non", "Non", "Oui"],
    ["Annuaire de participants & invitations tracées", "Partiel", "Non", "Oui"],
    ["Accès multi-rôles (admin/jury/porteur)", "Partiel", "Partiel", "Oui"],
], widths=[4.6, 3.9, 3.6, 3.9], cap_text="Comparaison des offres tunisiennes concurrentes", font_size=9.5)

H3("1.3.2 Critique de l'existant")
P("Cette analyse de l'existant fait ressortir quelques faiblesses récurrentes :")
bullet("**Outillage fragmenté.** Les candidatures, la grille de sélection, la notation du jury, "
       "l'agenda et les communications vivent chacun dans leur coin — fichiers séparés, boîtes de "
       "réception différentes. Rien n'est centralisé, et les mêmes données sont recopiées d'un "
       "outil à l'autre.")
bullet("**Aucune source unique de vérité pour le parcours.** Le calendrier du programme — ses "
       "sessions, leur ordre, leur imbrication (une journée de formation glissée dans une période "
       "d'incubation), leurs dépendances — n'est modélisé nulle part. En pratique, il n'existe que "
       "dans la tête du responsable de programme et dans quelques documents épars.")
bullet("**Évaluation manuelle et non reproductible.** Les jurys sont contactés par courriel, "
       "reçoivent les candidatures en pièce jointe, puis les notent sur papier ou dans des "
       "tableurs improvisés. Consolider des scores pondérés par critère, par jury et par session "
       "devient vite fastidieux — et réévaluer le même candidat lors d'une session ultérieure, un "
       "pitch day après une présélection par exemple, l'est plus encore.")
bullet("**Contrôle d'accès faible et informel.** Certaines sessions sont purement internes — "
       "planification, délibération du jury, recalibrage des scores — et ne devraient jamais être "
       "visibles des startups ni du public. Or rien ne l'empêche structurellement. À l'inverse, "
       "un jury externe invité à une seule session privée ne devrait pas, en théorie, accéder au "
       "reste du programme.")
bullet("**Traçabilité insuffisante des communications.** Lorsqu'un responsable envoie une "
       "invitation ou une mise à jour, il n'existe aucun enregistrement fiable de qui a été "
       "contacté, quand, avec quel message, ni même si le courriel a bien atteint son destinataire.")
bullet("**Absence de métriques consolidées.** Les métriques d'entonnoir (candidatures reçues, "
       "évaluées, acceptées), l'avancement des jurys, la répartition par secteur ou le classement "
       "des meilleurs candidats — tout cela doit être recalculé à la main pour chaque rapport.")
P("Additionnées, ces faiblesses ont un coût réel : du **temps perdu** en tâches manuelles "
  "répétitives, des **erreurs** de transcription et de consolidation des scores, des **retards** "
  "de décision, une certaine **frustration** des jurys et des porteurs face à un processus "
  "qu'ils ne maîtrisent pas vraiment, et beaucoup de **difficulté** à produire le reporting "
  "transparent qu'attendent financeurs et partenaires.")

H2("1.4 Solution proposée")
P("À partir de ces constats, la solution proposée prend la forme d'une **plateforme web unique "
  "et sécurisée**, capable de :")
bullet("modéliser un programme comme un parcours structuré et visuel de sessions typées ;")
bullet("recevoir et gérer les candidatures en ligne ;")
bullet("permettre une évaluation par jury reproductible et par session, assistée le cas échéant "
       "par l'intelligence artificielle ;")
bullet("gérer chaque participant et chaque organisation, et communiquer avec eux de façon traçable ;")
bullet("appliquer strictement les règles métier — validation des dates, visibilité des sessions — "
       "aussi bien côté serveur que dans l'interface ;")
bullet("distinguer clairement les rôles : administrateur, responsable de programme, porteur de "
       "projet, jury, mentor et grand public.")
P("Parmi ces piliers, l'un concentre l'essentiel de l'originalité du projet et mérite d'être "
  "énoncé clairement. La plateforme ne se contente pas de numériser la journée de présentation ; "
  "elle rend l'**évaluation d'un pitch mesurable**. Là où un incubateur s'appuie aujourd'hui sur "
  "l'impression subjective d'un jury, notée après coup, la plateforme produit des preuves : à "
  "quelle vitesse le porteur a réellement parlé, avec quelle intensité et quelle variation de "
  "voix, où sont tombés les vrais silences, combien d'hésitations ont échappé et exactement "
  "quand. Ce sont ces mesures — et non une impression — qui fondent la notation, et chaque "
  "remarque rendue au porteur est ancrée à un horodatage qu'il peut rejouer. Cela transforme un "
  "verdict ponctuel en quelque chose contre quoi un porteur peut s'entraîner, ce qui explique "
  "que le modèle distingue les pitchs d'entraînement, dont la progression est suivie d'une "
  "tentative à l'autre, du pitch final qui compte.")
bullet("transformer la journée de présentation en outil de coaching : transcrire chaque pitch "
       "automatiquement, mesurer objectivement la diction à partir du signal audio, et retourner "
       "une analyse notée, horodatée et actionnable ;")
bullet("exécuter toute cette analyse sur des modèles auto-hébergés, de sorte qu'aucune vidéo de "
       "candidat ne quitte l'infrastructure.")
P("Les chapitres suivants parcourent chacun de ces piliers à leur tour : le constructeur de "
  "parcours visuel, la gestion des sessions typées et contrôlées en visibilité, les candidatures "
  "en ligne, l'évaluation assistée par IA (surnommée « **Medi** »), la gestion des participants "
  "avec invitations tracées, et les catalogues de référence administrables.")

H2("1.5 Méthodologie de travail")
P("Le choix de la méthodologie de développement compte autant que le reste, surtout pour un "
  "projet dont les besoins évoluent en cours de route. Avant de justifier le choix retenu, cette "
  "section présente brièvement deux approches de référence.")
H3("1.5.1 Le cycle en cascade")
P("Le modèle **en cascade** progresse séquentiellement : chaque phase — spécification, "
  "conception, implémentation, tests, déploiement — doit être terminée avant de passer à la "
  "suivante. Lorsque les besoins sont figés dès le départ, cette prévisibilité est un atout réel. "
  "Mais dès que les besoins se précisent en cours de route, le modèle montre ses limites, car "
  "revenir en arrière devient coûteux.")
H3("1.5.2 Les méthodes agiles (Scrum)")
P("Les méthodes **agiles** — **Scrum** étant la plus connue — fonctionnent tout autrement : le "
  "développement est découpé en **itérations courtes**, les sprints, chacune livrant un "
  "incrément fonctionnel revu avec les parties prenantes avant de passer à la suivante. "
  "L'adaptation aux besoins devient continue, les problèmes remontent plus tôt, et la "
  "collaboration entre développeurs et parties prenantes métier se resserre.")
H3("1.5.3 Méthodologie retenue")
P("Pour ce projet, **les besoins n'étaient pas entièrement connus au départ** — ils se sont "
  "précisés à mesure que le travail avançait. Le modèle de session, les règles de visibilité, le "
  "flux d'évaluation : chacun a mûri sur plusieurs itérations, en discussion avec l'organisme "
  "d'accueil. Un cycle en cascade, qui suppose tout figé dès le premier jour, n'aurait simplement "
  "pas convenu à cette réalité. C'est pourquoi une approche **agile, itérative et incrémentale** "
  "a été retenue : le travail avance par incréments courts, chacun livrant une tranche verticale "
  "cohérente et déployable — back-end et front-end ensemble — intégrée en continu et validée "
  "avant de passer à la suite.")
P("**UML**, de son côté, est utilisé tout au long du projet comme **aide à la conception et à la "
  "communication** : il aide à clarifier les acteurs, les cas d'utilisation, le modèle du domaine "
  "et les principaux flux dynamiques, sans jamais chercher à figer une spécification exhaustive "
  "en amont. Cette agilité légère, guidée par les modèles, est proche dans l'esprit du "
  "**Two-Track Unified Process (2TUP)** et des pratiques **Scrum**.")
P("Le tableau ci-dessous résume cette comparaison.")
table([
    ["Critère", "Cascade", "Scrum (retenu)"],
    ["Adaptation aux besoins évolutifs", "Faible", "Élevée"],
    ["Livraison de valeur", "En fin de projet", "À chaque itération"],
    ["Risque de reprise coûteuse", "Élevé", "Faible"],
    ["Adapté à des besoins encore incertains", "Non", "Oui"],
], widths=[6.5, 4.75, 4.75], cap_text="Comparaison des méthodologies de développement")
H3("1.5.4 Le cadre Scrum : rôles, artefacts et événements")
P("Scrum organise le travail autour de trois **rôles**, trois **artefacts** et un ensemble "
  "d'**événements**. Les trois rôles sont le *Product Owner*, garant de la valeur et de la "
  "priorisation du backlog ; le *Scrum Master*, garant du cadre et de la levée des obstacles ; et "
  "l'*équipe de développement*, qui conçoit et livre l'incrément. Les trois artefacts sont le "
  "**backlog produit** (la liste priorisée de tout ce que le produit doit faire, exprimée en "
  "*user stories*), le **backlog de sprint** (le sous-ensemble sélectionné pour l'itération "
  "courante) et l'**incrément** (la tranche livrable produite à la fin du sprint). Les événements "
  "rythment l'itération : la *planification de sprint* (que livrer et comment), la *mêlée "
  "quotidienne* (synchronisation courte), la *revue de sprint* (démonstration de l'incrément aux "
  "parties prenantes) et la *rétrospective* (amélioration de la façon de travailler).")
P("Le tableau ci-dessous précise comment ces rôles se sont incarnés dans le contexte, réduit, "
  "de ce projet de fin d'études.")
table([
    ["Rôle Scrum", "Incarnation dans le projet"],
    ["Product Owner", "L'encadrant entreprise (Yassine Mechich), qui priorise les besoins et valide les incréments au regard des attentes de Medianet."],
    ["Scrum Master", "Rôle partagé, assurant le respect du cadre itératif et la levée des obstacles techniques."],
    ["Équipe de développement", "L'étudiant (Yassine Cherni), en charge de la conception, du développement full-stack et des tests de chaque incrément."],
], widths=[4.5, 11.5], cap_text="Rôles Scrum et leur incarnation dans le projet")
P("Concrètement, l'approche adoptée repose sur les points suivants :")
bullet("**Backlog produit.** Les besoins sont exprimés en **user stories** par acteur, puis priorisés.")
bullet("**Incréments / sprints.** Chaque incrément vise un ensemble cohérent de user stories — "
       "d'abord le constructeur de programme, puis les candidatures, l'évaluation, la notation IA, "
       "la visibilité, et enfin les participants.")
bullet("**Intégration et livraison continues.** Chaque changement passe par le typage statique "
       "(`tsc`), une compilation (`mvn` ou build Next.js) et un déploiement en conteneurs "
       "(`docker compose build` / `up`).")
bullet("**Isolation en microservices.** Chaque contexte métier est développé et déployé "
       "indépendamment, avec sa propre base de données.")
bullet("**Vérification.** Des scénarios de test manuels et des vérifications de bout en bout "
       "accompagnent chaque incrément, couvrant systématiquement chaque type de session et chaque rôle.")
H2("Conclusion")
P("Voilà donc le décor : un organisme d'accueil, un problème qui s'est avéré réel plutôt "
  "qu'inventé, une première esquisse de la plateforme, et la méthodologie Scrum qui a guidé sa "
  "construction. Avant de dérouler les sprints, le chapitre suivant établit les fondements "
  "théoriques sur lesquels repose la partie la plus originale du projet : l'analyse objective du pitch.")
pagebreak()

# =========================================================================
# CHAPITRE 2 — FONDEMENTS THÉORIQUES
# =========================================================================
H1("Chapitre 2 — État de l'art et fondements théoriques")
H2("Introduction")
P("La plateforme décrite dans ce rapport ne se contente pas de stocker et d'afficher des "
  "données : elle mesure une performance humaine et émet un jugement à son sujet. Cette ambition "
  "ne peut reposer sur l'intuition. Ce chapitre établit donc le socle scientifique sur lequel la "
  "solution est construite — comment la parole est transcrite automatiquement, comment les "
  "qualités d'une voix peuvent être quantifiées à partir du signal lui-même, ce qu'un modèle de "
  "vision peut ou ne peut pas voir de façon fiable, et sous quelles conditions un modèle de "
  "langage peut raisonnablement servir d'évaluateur. Chaque notion introduite ici est réutilisée "
  "plus loin : le chapitre 4 conçoit autour d'elle, le chapitre 5 l'implémente, et le chapitre 7 "
  "mesure si elle a tenu.")
H2("2.1 Reconnaissance automatique de la parole")
P("L'analyse d'un pitch commence par le transformer en texte exploitable. Cette section rappelle "
  "le fonctionnement de la reconnaissance automatique de la parole, du modèle historique en "
  "composants séparés aux architectures neuronales de bout en bout, puis met en évidence un biais "
  "de ces systèmes qui se révèle déterminant pour ce projet.")
H3("2.1.1 Des modèles acoustiques aux architectures de bout en bout")
P("La reconnaissance automatique de la parole (ASR, *Automatic Speech Recognition*) convertit un "
  "signal audio en texte. Historiquement, la tâche était scindée en composants séparés — un "
  "modèle acoustique reliant le son aux phonèmes, un lexique de prononciation, et un modèle de "
  "langage réordonnant les candidats en phrases plausibles. Chaque composant était entraîné à "
  "part, et les erreurs s'accumulaient le long de la chaîne. Les systèmes modernes remplacent ce "
  "pipeline par un unique réseau de neurones **de bout en bout** entraîné à faire correspondre "
  "directement l'audio à des caractères, ce qui supprime le lexique construit à la main et laisse "
  "le modèle apprendre seul la ponctuation et la casse.")
P("**Whisper**, publié par OpenAI en 2022, appartient à cette famille. C'est un Transformer "
  "encodeur–décodeur : l'encodeur consomme un spectrogramme log-Mel — une représentation "
  "temps–fréquence du son — et le décodeur génère le texte jeton par jeton, chaque jeton "
  "conditionné par ceux déjà produits. Sa particularité ne tient pas à son architecture mais à "
  "son régime d'entraînement : il a été entraîné sur environ 680 000 heures d'audio multilingue "
  "faiblement supervisé collecté sur le web, ce qui le rend robuste aux accents, au bruit de fond "
  "et aux conditions d'enregistrement sans aucun réglage fin. Cette robustesse est la raison de "
  "son choix ici : un pitch est filmé dans une salle, pas dans un studio.")
H3("2.1.2 Le biais de normalisation et sa conséquence pour ce projet")
P("Ce même régime d'entraînement porte un biais qui se révèle décisif pour ce travail. Whisper a "
  "été entraîné à produire une prose propre et lisible — le genre de texte que l'on trouve dans "
  "les sous-titres et les transcriptions — et il apprend donc à omettre ce qu'un transcripteur "
  "humain omettrait : les hésitations. Les pauses remplies (« euh », « hum »), les faux départs "
  "et les répétitions sont systématiquement absents de sa sortie. Pour une tâche de "
  "sous-titrage, c'est une qualité. Pour une application dont le but affiché est de détecter "
  "l'hésitation, cela détruit le signal même que l'on cherche.")
P("Deux propriétés du décodeur permettent de contourner cela. D'abord, le décodeur est "
  "autorégressif et conditionné par son propre contexte : un préfixe qu'on lui fournit "
  "(l'*initial prompt*) influence le registre de tout ce qui suit. Ensuite, ce préfixe n'est pas "
  "transcrit ; il ne fait que conditionner le style. Amorcer le décodeur avec un échantillon "
  "délibérément disfluent le pousse donc vers une transcription verbatim. Le chapitre 5 détaille "
  "cette technique et le chapitre 7 en mesure l'effet.")
P("Une seconde propriété est exploitée plus loin : les **horodatages au niveau du mot**. En "
  "alignant l'attention croisée du décodeur sur les trames d'entrée, Whisper peut retourner un "
  "instant de début et de fin pour chaque mot, ainsi que la probabilité qu'il lui a attribuée. "
  "Ces trois quantités — le mot, sa durée et sa confiance — sont la matière première des mesures "
  "de disfluence définies en section 2.2.3.")
H2("2.2 Mesure objective du signal de parole")
P("La transcription dit *ce qui* a été dit. Elle ne dit rien de *comment*. Tout ce qui touche à "
  "la diction — si la voix porte, si elle est monotone, si l'orateur hésite — doit être mesuré "
  "sur la forme d'onde, indépendamment de tout modèle. Cette section définit les quantités "
  "utilisées.")
H3("2.2.1 L'intensité sonore : la norme EBU R 128")
P("L'intensité sonore perçue n'est pas l'amplitude physique d'un signal : l'oreille n'est pas "
  "également sensible à toutes les fréquences. La recommandation **EBU R 128**, avec la norme "
  "ITU-R BS.1770, définit une mesure qui approche la perception. Le signal est d'abord filtré par "
  "une courbe de « pondération K », qui atténue les basses fréquences et accentue légèrement les "
  "hautes pour imiter la réponse de l'oreille ; la moyenne quadratique du signal filtré est "
  "ensuite intégrée dans le temps et exprimée en **LUFS** (*Loudness Units relative to Full "
  "Scale*), une échelle logarithmique où 0 LUFS est le maximum et où les valeurs sont donc "
  "négatives.")
P("Deux grandeurs dérivées de cette norme sont utilisées dans ce projet. L'**intensité intégrée** "
  "résume le niveau moyen de tout l'enregistrement : la pratique broadcast vise environ "
  "−23 LUFS, et un pitch mesuré bien en deçà — autour de −30 LUFS — est objectivement trop "
  "faible pour être entendu confortablement. La **plage d'intensité** (LRA), exprimée en LU, "
  "mesure l'écart entre les passages doux et forts. Une LRA basse est la signature d'une diction "
  "monotone ; un orateur qui varie l'intensité pour souligner un point en produit une plus "
  "élevée. Ces deux nombres remplacent le jugement subjectif « il a parlé trop doucement et "
  "d'une voix plate » par une mesure reproductible.")
H3("2.2.2 La détection des silences et son interaction avec l'ASR")
P("Une pause est un silence plus long qu'un seuil donné. En détecter une requiert deux "
  "paramètres : un niveau en dessous duquel le signal est considéré silencieux (en dB) et une "
  "durée minimale en dessous de laquelle le vide est une simple rupture articulatoire plutôt "
  "qu'une pause. Des intervalles détectés découlent le nombre de pauses, la plus longue, et le "
  "**ratio de parole** — la part de l'enregistrement qui est réellement de la parole.")
P("Un point méthodologique mérite d'être souligné, car il a causé une véritable erreur dans ce "
  "projet. La plupart des systèmes ASR, Whisper compris, appliquent un filtre de détection "
  "d'activité vocale (*VAD*) qui retire le silence avant transcription. En conséquence, les vides "
  "entre les segments transcrits ne sont pas les silences de l'enregistrement d'origine — ils "
  "ont déjà été éliminés. Dériver les pauses à partir des segments de transcription ne donne donc "
  "rien. Les silences doivent être mesurés sur l'audio brut, indépendamment de l'ASR. C'est un "
  "principe général : **une mesure doit être prise sur la source, jamais sur un artefact déjà "
  "transformé**.")
H3("2.2.3 Alignement au niveau du mot et marqueurs de disfluence")
P("La psycholinguistique distingue plusieurs marqueurs de disfluence : les pauses remplies "
  "(« euh », « hum »), les prolongations (un son tenu anormalement longtemps), les répétitions "
  "et réparations (faux départs), et les expressions d'atténuation (« je pense », « peut-être ») "
  "qui affaiblissent l'assertion. L'*uptalk* — une affirmation prononcée avec l'intonation d'une "
  "question — appartient à la même famille. Ces marqueurs sont étudiés comme indices de la charge "
  "cognitive et de la confiance de l'orateur, précisément ce qu'un jury perçoit sans pouvoir le "
  "quantifier.")
P("Les horodatages au niveau du mot en rendent plusieurs mesurables. Une prolongation, par "
  "exemple, peut se définir comme un mot court dont la durée dépasse un seuil. Cette définition "
  "n'est toutefois valide que si l'on corrige une propriété de l'aligneur : la fin d'un mot est "
  "placée là où le suivant commence, de sorte que tout silence qui suit un mot lui est imputé. Un "
  "mot suivi d'une pause paraît donc étiré. Soustraire à chaque mot les silences mesurés en "
  "section 2.2.2 restitue la durée voisée, qui est la quantité réellement intéressante. Le "
  "chapitre 7 quantifie l'écart entre la mesure naïve et la mesure corrigée.")
H2("2.3 Modèles vision–langage et limites de la preuve visuelle")
P("La présence physique — posture, regard, geste — fait partie de ce qu'un jury juge. Les "
  "**modèles vision–langage** la rendent partiellement accessible : un encodeur d'image projette "
  "une photo dans le même espace de représentation qu'un modèle de langage, qui peut alors "
  "décrire la scène en langage naturel. Un modèle compact de cette famille, exécuté localement, "
  "est utilisé ici pour décrire des trames échantillonnées du pitch.")
P("Trois limites régissent la façon dont une telle description peut légitimement être utilisée. "
  "D'abord, une trame échantillonnée est un instant, pas un comportement : une image toutes les "
  "quarante secondes montre des postures, pas des habitudes, et un trait observé une seule fois "
  "peut être un accident de timing. Ensuite, le décodage est stochastique : à température non "
  "nulle, le modèle échantillonne dans une distribution de probabilité, si bien que la même image "
  "peut produire des descriptions différentes d'une exécution à l'autre. La reproductibilité "
  "exige donc de fixer la température et la graine aléatoire. Enfin, ces modèles sont sujets à "
  "l'**hallucination** — produire une description plausible mais non étayée par les pixels — un "
  "risque qui croît quand le sujet occupe une petite part de l'image, ce qui est exactement le "
  "cas d'un orateur filmé sur une scène. Ces trois limites justifient la règle de conception "
  "adoptée au chapitre 4 : les observations visuelles sont traitées comme des **indices faibles**, "
  "et un trait vu dans une seule trame ne peut jamais, à lui seul, justifier un conseil.")
H2("2.4 Les grands modèles de langage comme évaluateurs")
P("Une fois la diction mesurée, reste à porter un jugement argumenté sur le pitch : c'est le rôle "
  "confié à un grand modèle de langage. Cette section en rappelle le principe et les biais "
  "documentés, puis énonce les deux garde-fous — ancrage et calibration — qui conditionnent son "
  "usage comme évaluateur.")
H3("2.4.1 Principe et biais connus")
P("Un grand modèle de langage (LLM) est un Transformer entraîné à prédire le jeton suivant sur "
  "de très grands corpus. Ses variantes instruites peuvent se voir demander d'appliquer une "
  "grille de notation à un texte, approche connue dans la littérature sous le nom de "
  "*LLM-as-a-judge*. Son attrait est évident — il produit un jugement argumenté et structuré en "
  "quelques secondes — tout comme ses faiblesses, documentées et toutes observées durant ce projet.")
P("La première est l'**hallucination** : le modèle énonce, avec fluidité et assurance, des faits "
  "qui ne sont pas dans son entrée. La deuxième est la **complaisance** (*sycophancy*) : une "
  "tendance à être agréable qui, sur une tâche d'évaluation, se traduit par une indulgence "
  "systématique. La troisième est la **sensibilité à la formulation du prompt** : une grille "
  "exprimée sous forme de plafonds rigides ne produit pas un correcteur plus strict mais un "
  "correcteur dégénéré, comme le montre expérimentalement le chapitre 7. La quatrième est le "
  "**non-déterminisme**, de même nature que celui du modèle de vision. Prises ensemble, ces biais "
  "interdisent de demander à un LLM de « noter ce pitch » et de prendre la réponse pour argent comptant.")
H3("2.4.2 Ancrage et calibration")
P("L'approche défendue dans ce rapport découle de ces limites et repose sur deux principes. "
  "L'**ancrage** (*grounding*) consiste à ne jamais demander au modèle un fait qu'il pourrait "
  "inventer : tout ce qui est factuel — débit, intensité, pauses, hésitations et leurs "
  "horodatages — est mesuré au préalable par les méthodes déterministes de la section 2.2, et le "
  "modèle n'est sollicité que pour interpréter ces mesures au regard des critères du programme. "
  "Son rôle passe de témoin à analyste. La **calibration** consiste à vérifier que l'échelle "
  "obtenue discrimine réellement : une grille qui attribue la même note à un excellent pitch et à "
  "un mauvais ne porte aucune information, quelle que soit la qualité de sa prose. La calibration "
  "est donc une question expérimentale, tranchée sur un corpus de référence, et le chapitre 7 la "
  "traite comme telle.")
P("Une précaution structurelle les complète : le score global n'est pas demandé au modèle mais "
  "dérivé de ses scores par dimension. Cela empêche une « impression » holistique de l'emporter "
  "sur le détail, et rend la note réconciliable avec le profil montré à l'utilisateur.")
H2("2.5 Fondements architecturaux et de sécurité")
P("Les besoins fonctionnels reposent aussi sur quelques notions établies, rappelées ici "
  "brièvement. Une architecture **microservices** décompose une application en services "
  "déployables indépendamment, chacun propriétaire de ses données ; elle convient à ce projet "
  "car le pipeline média a des besoins en ressources — CPU, modèles, ffmpeg — qui n'ont rien de "
  "commun avec ceux d'un service CRUD, et qui doivent pouvoir tomber ou être redéployés sans "
  "faire chuter la plateforme. Un registre de services et une **passerelle d'API** fournissent "
  "respectivement la découverte et un point d'entrée unique.")
P("La sécurité repose sur le **JSON Web Token** (JWT, RFC 7519), un jeton signé et autonome qui "
  "porte l'identité et les droits de l'utilisateur, permettant à chaque service d'autoriser une "
  "requête sans session partagée. Il est combiné au **contrôle d'accès basé sur les rôles** "
  "(RBAC), étendu ici par des permissions par utilisateur et l'héritage de rôles. Enfin, les "
  "**Server-Sent Events** (SSE) fournissent un canal unidirectionnel serveur→client sur HTTP, "
  "utilisé à deux fins : propager un changement de droits à un utilisateur connecté, et diffuser "
  "l'avancement du pipeline d'analyse, dont la durée rend une requête synchrone inacceptable.")
H2("Conclusion")
P("Ce chapitre a établi ce sur quoi s'appuie le reste du rapport : un système ASR robuste mais "
  "biaisé vers la prose propre, dont il faut contourner le biais pour observer l'hésitation ; un "
  "ensemble de mesures acoustiques normalisées — intensité, plage d'intensité, silences — qui "
  "objectivent la diction ; l'alignement au niveau du mot, utilisable seulement une fois corrigé "
  "du silence imputé à chaque mot ; un modèle de vision dont la sortie est un indice faible et "
  "stochastique ; et un modèle de langage qui n'est un évaluateur crédible que s'il est ancré "
  "dans des mesures préalables et calibré sur un corpus de référence. Ces principes ne sont pas "
  "décoratifs : chacun correspond à une erreur réellement commise et corrigée durant le projet, "
  "comme le documente le chapitre 7. Le chapitre suivant ouvre le *Sprint 0* et transforme la "
  "vision en besoins.")
pagebreak()

# =========================================================================
# CHAPITRE 3 — SPRINT 0 / ÉTUDE PRÉLIMINAIRE
# =========================================================================
H1("Chapitre 3 — Étude préliminaire (Sprint 0)")
H2("Introduction")
P("Le premier sprint d'un projet Scrum, le *Sprint 0*, ne livre pas de fonctionnalité : il "
  "prépare le terrain. Ce chapitre en tient le rôle. Il transforme la vision en besoins — "
  "fonctionnels et non fonctionnels — identifie les acteurs et les cas d'utilisation, puis "
  "constitue le **backlog produit** sous forme de *user stories* priorisées. Il fixe ensuite la "
  "**feuille de route** en releases, arrête l'**architecture globale** et met en place "
  "l'**environnement de travail**. À l'issue de ce sprint, l'équipe dispose de tout ce qu'il faut "
  "pour commencer à livrer.")
H2("3.1 Spécification des besoins")
P("La vision derrière **Medianet Incubateur** tient en quelques lignes :")
P("« Offrir aux incubateurs une plateforme unique qui transforme un programme en un **parcours "
  "structuré et visuel** de sessions typées et contrôlées en accès ; permet aux candidats de "
  "**postuler en ligne** ; permet aux jurys d'**évaluer** (assistés par l'intelligence "
  "artificielle) ; et tient chaque participant **informé** par des communications traçables — le "
  "tout en garantissant une **validation stricte des données** et une **séparation claire des "
  "rôles**. »", indent=False)
H3("3.1.1 Besoins fonctionnels")
P("Pour plus de clarté, les besoins fonctionnels sont regroupés par module.")
P("**M1 — Authentification et comptes.** Ce module est la fondation sur laquelle repose le reste "
  "de la plateforme : il établit qui est un utilisateur et ce qu'il a le droit de faire. Plutôt "
  "que de lier les droits à une poignée de rôles figés, Medianet superpose aux rôles des "
  "permissions fines par utilisateur, de sorte qu'un administrateur peut accorder à un seul "
  "mentor l'accès à un écran sans inventer un nouveau rôle.")
bullet("Inscription et connexion avec **JWT** ; réinitialisation du mot de passe.")
bullet("Gestion du profil utilisateur, différenciée par rôle (porteur, jury, mentor).")
bullet("Administration des utilisateurs, de leurs **rôles** et de leurs **permissions fines**.")
P("**M2 — Programmes et parcours.** M2 est le cœur du produit. Un programme n'est pas seulement "
  "une fiche avec des dates : il porte un parcours visuel fait de sessions typées — candidature, "
  "présélection, journée de pitch, incubation, démo day — que l'on peut disposer sur une frise et "
  "imbriquer l'une dans l'autre (une session-jour dans une plage). Comme le statut du programme "
  "est dérivé automatiquement de ces sessions et de leurs dates, l'équipe est dispensée de la "
  "tenue manuelle qui laisse habituellement un programme dériver par rapport à la réalité.")
bullet("Créer, consulter, modifier et supprimer des programmes, avec leurs métadonnées et un "
       "formulaire de candidature personnalisable.")
bullet("Gérer les **critères d'évaluation** (nom, poids, note maximale) par programme.")
bullet("Construire le **parcours visuel** : créer des sessions (jour ou plage), imbriquer les "
       "sessions-jour, les organiser en couloirs, définir leur type, leur visibilité, leurs dates, "
       "leur lieu, leurs propriétaires et invités, et fixer les poids de critères par session.")
bullet("Éditer un **agenda d'activités** (grille horaire) pour les sessions-jour ; réutiliser "
       "des *presets* et des modèles de parcours complets.")
bullet("**Synchroniser automatiquement le statut du programme** d'après les dates et statuts de ses sessions.")
P("**M3 — Candidatures.** Ce module régit la façon dont les porteurs entrent dans un programme. "
  "Chaque programme définit son propre formulaire de candidature, de sorte qu'un incubateur "
  "collecte exactement l'information qui l'intéresse, et chaque candidature suit ensuite un cycle "
  "de statut clair, du dépôt à l'acceptation ou au rejet.")
bullet("Déposer une candidature à un programme ouvert via son propre formulaire, en liant "
       "l'**organisation** et l'**équipe** du porteur.")
bullet("Gérer le **cycle de vie du statut** d'une candidature (en attente, en évaluation, acceptée, rejetée).")
bullet("Consulter le **détail complet** d'une candidature, y compris les réponses soumises.")
P("**M4 — Évaluation.** L'évaluation est le moment où un programme prend réellement ses "
  "décisions, et elle combine deux voix. Les jurys notent une candidature au regard de critères "
  "pondérés ; à leurs côtés, **Medi** — l'évaluateur IA — lit la même candidature et retourne une "
  "appréciation par critère assortie d'un score pondéré justifié. Medi est délibérément positionné "
  "comme un second avis cohérent qui ancre chaque note dans les réponses soumises, non comme un "
  "remplacement du jugement humain, ce qui rend son score défendable devant un jury.")
bullet("**Affecter des jurys** à une session d'évaluation.")
bullet("**Noter** une candidature au regard de critères pondérés ; autoriser la **réévaluation "
       "par session** (un candidat est noté à neuf dans chaque session d'évaluation).")
bullet("Lancer un **score IA (Medi)** produisant des notes par critère, un score pondéré, une "
       "recommandation et une synthèse ; permettre aussi la notation par les administrateurs.")
bullet("Fournir un **tableau de bord consolidé** avec classement et export.")
P("**M5 — Participants et communications.** M5 répond à deux questions quotidiennes — qui est "
  "impliqué dans un programme, et comment les joindre. Les personnes peuvent être ajoutées depuis "
  "des comptes existants ou invitées par courriel, et chaque message, modèle personnalisé ou "
  "notification par session, est archivé et porte un jeton de suivi. La communication devient "
  "donc auditable.")
bullet("Lister les **personnes associées** à un programme (organisations, membres, porteurs, jurys).")
bullet("**Ajouter des utilisateurs déjà inscrits** ou **inviter des personnes par courriel**.")
bullet("Composer une **invitation personnalisée** à partir d'un modèle avec variables et **lien "
       "de suivi** ; **prévisualiser avant l'envoi**.")
bullet("Envoyer des **notifications par session** et conserver une **archive complète** de chaque envoi.")
P("**M6 — Visibilité et contrôle d'accès.** Ce module transversal décide de ce que chaque "
  "utilisateur a même le droit de voir. La visibilité est une propriété de la session — Visible, "
  "Masquée ou Privée — et elle est appliquée côté serveur d'après le rôle et les invitations du "
  "demandeur, jamais simplement cachée dans l'interface.")
bullet("Définir la **visibilité** d'une session : **Visible**, **Masquée (interne)** ou **Privée**.")
bullet("**Appliquer** la visibilité côté serveur selon le **rôle** et les **invitations** du "
       "demandeur, et la refléter dans l'interface.")
P("**M7 — Catalogues de référence.** M7 maintient la cohérence du vocabulaire de la plateforme. "
  "Types d'organisation, secteurs de programme et types de session sont gérés comme des "
  "catalogues éditables par l'administrateur.")
bullet("Gérer les **catalogues** (types d'organisation, secteurs de programme, types de session) "
       "utilisés par les formulaires.")
P("**M8 — Journée de présentation et analyse IA du pitch.** M8 est ce qui distingue le plus "
  "clairement Medianet d'un outil de gestion ordinaire. Toute session peut être marquée pour "
  "analyse vidéo, après quoi le pitch d'un porteur est transcrit et horodaté, et sa diction "
  "mesurée objectivement à partir du signal lui-même — débit, intensité et sa dynamique, pauses, "
  "et habitudes inconscientes mot à mot telles que sons de remplissage, mots étirés et "
  "hésitations. La présence physique est décrite à partir de trames échantillonnées par un modèle "
  "de vision, et le pitch est ensuite noté par dimension au regard des critères du programme, "
  "avec des moments-clés horodatés et un plan de coaching chiffré. Fait crucial, tout ce pipeline "
  "s'exécute sur l'infrastructure de l'incubateur : la vidéo ne la quitte jamais.")
bullet("Marquer n'importe quelle session pour analyse vidéo IA d'une simple case, "
       "indépendamment du type de session.")
bullet("Laisser un porteur téléverser plusieurs vidéos d'entraînement et une vidéo de pitch "
       "final (jusqu'à 2 Go chacune), stockées dans MinIO.")
bullet("Transcrire le pitch automatiquement et l'horodater, en conservant les vraies hésitations "
       "plutôt qu'un texte nettoyé.")
bullet("Mesurer la diction objectivement depuis la forme d'onde : débit, intensité, plage "
       "d'intensité, vraies pauses et ratio de parole.")
bullet("Détecter les habitudes de parole inconscientes mot à mot : sons de remplissage, sons "
       "étirés, répétitions, atténuations et questions-tag.")
bullet("Décrire la présence physique à partir de trames échantillonnées via un modèle de vision "
       "exécuté localement, traité comme indice faible.")
bullet("Noter le pitch par dimension au regard des critères du programme, après détection de son "
       "format (démo day, compétition, investisseur, interne).")
bullet("Retourner des moments-clés horodatés, une cartographie des critères et un plan de "
       "coaching chiffré (impact en points, effort, formulation prête à dire).")
bullet("Diffuser en direct les étapes du pipeline vers l'interface, avec des estimations de temps "
       "apprises de vraies exécutions.")
bullet("Afficher des analyses de progression à trois niveaux : global, par programme et par phase d'entraînement.")
H3("3.1.2 Besoins non fonctionnels")
table([
    ["Qualité", "Exigence"],
    ["Sécurité", "Authentification **JWT** sans état à la passerelle et dans chaque service ; autorisation par **rôle** et **permission** (`module:action`) ; filtrage des données par rôle, notamment pour la visibilité des sessions."],
    ["Scalabilité", "Microservices indépendants et sans état derrière une passerelle avec découverte de services ; **une base de données par service**."],
    ["Maintenabilité & extensibilité", "Contextes métier bien délimités ; couches DTO ; une **couche de validation centralisée** avec codes d'erreur explicites ; un **modèle de session extensible** (nouveaux types sans refonte du schéma)."],
    ["Utilisabilité", "Deux interfaces dédiées ; une interface de session **à onglets** lisible ; **confirmation avant ajout** et **prévisualisation avant modification** pour les actions sensibles."],
    ["Performance", "Chargement paresseux, requêtes **parallélisées**, pas d'appels **N+1** inter-services grâce aux champs d'affichage dénormalisés."],
    ["Fiabilité", "Appels inter-services **résilients** (dégradation gracieuse), gestion structurée des erreurs, déploiements reproductibles."],
    ["Portabilité", "Solution entièrement **conteneurisée** (Docker), configuration pilotée par variables d'environnement."],
    ["Localisation", "Interface en langue française, adaptée au contexte local."],
], widths=[3.6, 12.4], font_size=9.5)
H2("3.2 Identification des acteurs")
P("Un **acteur** est un rôle joué par un utilisateur, ou par un système externe, qui interagit "
  "avec la plateforme.")
table([
    ["Acteur", "Type", "Description"],
    ["Administrateur", "Humain (principal)", "Contrôle total sur les programmes, parcours, critères, candidatures, évaluations, utilisateurs, rôles, catalogues et communications."],
    ["Responsable de programme", "Humain", "Un administrateur restreint, cadré par des permissions (`programmes:*`, `candidatures:*`, …) aux opérations d'un programme."],
    ["Porteur de projet", "Humain (principal)", "Fondateur qui postule, gère son organisation et suit le parcours et ses tâches."],
    ["Jury", "Humain (principal)", "Évalue les candidatures qui lui sont affectées, par session, éventuellement assisté par l'IA."],
    ["Mentor", "Humain", "Accompagne les startups sélectionnées ; consulte les sessions et tâches qui lui sont assignées."],
    ["Membre d'organisation", "Humain", "Membre d'une organisation participante ; suit les mises à jour du programme."],
    ["Visiteur", "Humain", "Utilisateur anonyme qui parcourt les programmes publics (sessions Visibles uniquement)."],
    ["Medi (LLM)", "Système externe", "Le service d'IA qui note les candidatures et analyse les pitchs."],
    ["Relais SMTP (Brevo)", "Système externe", "Achemine les courriels d'invitation et de notification."],
], widths=[3.4, 3.0, 9.6], cap_text="Identification des acteurs", font_size=9.5)
H2("3.3 Diagramme de cas d'utilisation global")
P("Le diagramme ci-dessous résume la façon dont ces acteurs interagissent avec la plateforme.")
figure('image4.png', "Diagramme de cas d'utilisation global")
P("La figure ci-dessus présente l'ensemble des cas d'utilisation. Dix d'entre eux ont été "
  "retenus pour une spécification détaillée ; ils sont répartis, dans les chapitres suivants, "
  "selon la release qui les livre.")
table([
    ["Code", "Cas d'utilisation", "Acteur principal", "Release"],
    ["UC-1", "S'authentifier", "Tous les acteurs inscrits", "R1"],
    ["UC-3", "Concevoir un programme et ses critères", "Administrateur", "R1"],
    ["UC-4", "Construire le parcours visuel", "Administrateur", "R1"],
    ["UC-5", "Définir le type et la visibilité d'une session", "Administrateur", "R1"],
    ["UC-2", "Déposer une candidature", "Porteur", "R2"],
    ["UC-6", "Affecter un jury", "Administrateur", "R2"],
    ["UC-7", "Évaluer une candidature (par session)", "Jury", "R2"],
    ["UC-8", "Noter une candidature par IA", "Jury / Administrateur", "R2"],
    ["UC-9", "Gérer les participants et envoyer des invitations", "Administrateur", "R3"],
    ["UC-10", "Suivre un programme", "Porteur / Jury", "R3"],
], widths=[1.4, 6.4, 5.2, 3.0], cap_text="Liste des cas d'utilisation majeurs et leur release", font_size=9.5)
H2("3.4 Backlog produit")
P("Le backlog produit exprime les besoins sous forme de *user stories* — « en tant que <rôle>, "
  "je veux <objectif> afin de <bénéfice> ». Chaque story porte une **priorité** (Haute, Moyenne, "
  "Basse) et est rattachée à une **release**. Le tableau ci-dessous en donne l'essentiel ; il "
  "constitue le point de départ ordonnancé de tous les sprints qui suivent.")
table([
    ["#", "En tant que…", "je veux…", "Prio.", "Rel."],
    ["US-01", "utilisateur", "m'inscrire et me connecter en sécurité (JWT)", "Haute", "R1"],
    ["US-02", "administrateur", "gérer les utilisateurs, rôles et permissions fines", "Haute", "R1"],
    ["US-03", "responsable", "créer et configurer un programme et ses critères", "Haute", "R1"],
    ["US-04", "responsable", "construire un parcours visuel de sessions typées et imbriquées", "Haute", "R1"],
    ["US-05", "responsable", "définir le type et la visibilité de chaque session", "Haute", "R1"],
    ["US-06", "responsable", "voir le statut du programme se synchroniser depuis ses sessions", "Moyenne", "R1"],
    ["US-07", "porteur", "postuler en ligne via le formulaire propre au programme", "Haute", "R2"],
    ["US-08", "porteur", "lier ou créer mon organisation et mon équipe", "Moyenne", "R2"],
    ["US-09", "administrateur", "affecter des jurys à une session d'évaluation", "Haute", "R2"],
    ["US-10", "jury", "noter une candidature par critères pondérés, par session", "Haute", "R2"],
    ["US-11", "jury", "obtenir un score IA (Medi) par critère, réutilisable", "Haute", "R2"],
    ["US-12", "porteur", "téléverser des pitchs (entraînement et final) jusqu'à 2 Go", "Haute", "R2"],
    ["US-13", "porteur", "recevoir une analyse mesurée et horodatée de mon pitch", "Haute", "R2"],
    ["US-14", "administrateur", "gérer les participants et envoyer des invitations tracées", "Haute", "R3"],
    ["US-15", "administrateur", "notifier les participants d'une session et archiver les envois", "Moyenne", "R3"],
    ["US-16", "porteur / jury", "suivre le parcours, mes tâches et mes mises à jour", "Moyenne", "R3"],
    ["US-17", "visiteur", "parcourir les programmes publics et leurs sessions visibles", "Basse", "R3"],
    ["US-18", "administrateur", "gérer les catalogues de référence", "Basse", "R3"],
], widths=[1.3, 2.7, 8.0, 1.9, 1.1], cap_text="Backlog produit — user stories priorisées", font_size=9.0)
H2("3.5 Planification des releases")
P("Les user stories du backlog sont regroupées en incréments — les **sprints** — eux-mêmes "
  "réunis en trois **releases** cohérentes livrées successivement. Le tableau ci-dessous retrace "
  "cette feuille de route, du Sprint 0 aux tests de bout en bout.")
table([
    ["Sprint", "Release", "Contenu"],
    ["Sprint 0", "—", "Montée en compétence, étude de l'existant, mise en place de l'environnement (Docker, dépôt Git, squelette microservices)."],
    ["Sprint 1", "R1", "Authentification JWT, gestion des utilisateurs et des rôles."],
    ["Sprint 2", "R1", "Modèle de programme et constructeur de parcours visuel (sessions jour/plage, imbrication)."],
    ["Sprint 3", "R2", "Candidatures en ligne et formulaire personnalisable."],
    ["Sprint 4", "R2", "Évaluation par jury, critères pondérés, réévaluation par session."],
    ["Sprint 5", "R2", "Intégration de la notation IA (Medi) et de l'analyse IA du pitch (transcription, mesures, vision)."],
    ["Sprint 6", "R3", "Visibilité des sessions (VISIBLE/MASQUÉE/PRIVÉE) et couche de validation centralisée."],
    ["Sprint 7", "R3", "Participants, invitations personnalisées et archive des communications."],
    ["Sprint 8", "—", "Tests de bout en bout, corrections, validation et rédaction du rapport."],
], widths=[2.2, 1.6, 12.2], cap_text="Feuille de route : sprints et releases", font_size=9.5)
H2("3.6 Architecture globale")
P("Les diagrammes de classes disent *ce que* le système sait ; le diagramme de composants dit "
  "*où* cette connaissance vit et *qui* a le droit de parler à qui. La figure ci-dessous se lit "
  "de haut en bas, dans l'ordre où une requête voyage réellement. Les deux front-ends Next.js ne "
  "s'adressent jamais directement à un service — ils ne connaissent que la **passerelle d'API**, "
  "unique point d'entrée public et unique endroit où un JWT est validé. Derrière elle se trouvent "
  "les services Spring Boot, chacun propriétaire d'un contexte borné et d'un schéma PostgreSQL, "
  "de sorte qu'aucun service ne lit les tables d'un autre ; **Eureka** assure la découverte.")
figure('image14.jpeg', "Diagramme de composants de la plateforme")
P("Le tableau ci-dessous précise la répartition des responsabilités entre les services.")
table([
    ["Service", "Port", "Responsabilité"],
    ["API Gateway", "8080", "Point d'entrée unique ; routage `/api/**` ; CORS ; transmission du JWT."],
    ["Eureka", "8761", "Registre / découverte de services."],
    ["auth-service", "8081", "Utilisateurs, rôles/permissions, organisations & membres, émission du JWT, téléversements (MinIO)."],
    ["programme-service", "8086", "Programmes, sessions (parcours), critères, partenaires, tâches, catalogues ; couche de validation ; filtrage de visibilité."],
    ["candidature-service", "8083", "Candidatures, affectations de jury, évaluations, listes restreintes."],
    ["notification-service", "8087", "Invitations, contacts/groupes, notifications de session, archive des courriels."],
    ["admin-ai-service", "8088", "Notation IA (Medi) et assistant d'administration."],
    ["pitch-media-service", "8089", "Pipeline vidéo : transcription (Whisper), mesures audio (ffmpeg), vision."],
], widths=[3.4, 1.3, 11.3], cap_text="Cartographie des services et des ports", font_size=9.5)
H2("3.7 Environnement de travail")
P("Le Sprint 0 met également en place l'environnement dans lequel tout le projet sera développé "
  "et exécuté. Cette section en récapitule les composantes logicielles et matérielles, ainsi que "
  "les frameworks et les langages retenus.")
H3("3.7.1 Environnement logiciel")
table([
    ["Catégorie", "Outils"],
    ["Système d'exploitation", "Windows 11 (développement), conteneurs Linux (exécution)"],
    ["Conteneurisation", "Docker, Docker Compose"],
    ["EDI", "IntelliJ IDEA, Visual Studio Code"],
    ["Outils API & BD", "Postman, pgAdmin / psql"],
    ["Gestion de versions", "Git"],
    ["Stockage objet", "MinIO (compatible S3)"],
    ["Courriel transactionnel", "Relais SMTP Brevo"],
    ["Intelligence artificielle", "OpenRouter / Ollama local (LLM), Whisper, modèle de vision"],
], widths=[4.5, 11.5], cap_text="Environnement logiciel", font_size=10)
H3("3.7.2 Environnement matériel")
table([
    ["Élément", "Spécification"],
    ["Machine de développement", "CPU multi-cœurs, RAM ≥ 16 Go recommandée, SSD"],
    ["Exécution", "Un hôte Docker faisant tourner une douzaine de conteneurs (services + bases + passerelle + deux front-ends + MinIO + runtime de modèles)"],
], widths=[4.5, 11.5], cap_text="Environnement matériel", font_size=10)
H3("3.7.3 Frameworks")
bullet("**Back-end :** Spring Boot, Spring Web, Spring Data JPA, Spring Security, Spring Cloud "
       "Gateway, Spring Cloud Netflix Eureka, Hibernate, Lombok.")
bullet("**Front-end :** Next.js (App Router), React, Tailwind CSS, Framer Motion, Recharts, "
       "lucide-react, react-hot-toast, Axios.")
bullet("**Média / IA :** ffmpeg, Whisper (transcription), Ollama (modèle de vision local).")
H3("3.7.4 Langages de programmation")
bullet("**Java 17** — les microservices.")
bullet("**TypeScript / JavaScript** — les front-ends Next.js.")
bullet("**Python** — le service de pipeline média (pitch-media-service).")
bullet("**SQL** — PostgreSQL.")
bullet("**HTML / CSS** — le balisage et le style Tailwind.")
H2("Conclusion")
P("À l'issue du Sprint 0, la vision est devenue un backlog produit priorisé, les acteurs et les "
  "cas d'utilisation sont posés, l'architecture microservices est arrêtée et l'environnement "
  "conteneurisé est opérationnel. La feuille de route découpe le travail en trois releases. Le "
  "chapitre suivant ouvre la première d'entre elles.")
pagebreak()

def uc_table(rows, cap_text):
    table([["Champ", "Description"]] + rows, widths=[3.2, 12.8], cap_text=cap_text, font_size=9.5)

# =========================================================================
# CHAPITRE 4 — RELEASE 1
# =========================================================================
H1("Chapitre 4 — Release 1 : Fondations et constructeur de parcours")
H2("Introduction")
P("La première release pose les fondations du produit et livre déjà son cœur métier. Elle "
  "couvre les sprints 1 et 2 : d'un côté l'**authentification** et la gestion des utilisateurs, "
  "des rôles et des permissions ; de l'autre le **modèle de programme** et le **constructeur de "
  "parcours visuel**, avec ses sessions typées, imbriquées et contrôlées en visibilité. À l'issue "
  "de cette release, un responsable peut se connecter, créer un programme et en dessiner le "
  "parcours complet. Ce chapitre suit le déroulé d'un sprint Scrum : backlog, analyse, "
  "conception, réalisation, puis revue et rétrospective.")
H2("4.1 Backlog de la release")
P("Cette release tire du backlog produit les user stories fondatrices, résumées ci-dessous.")
table([
    ["#", "User story", "Prio."],
    ["US-01", "En tant qu'utilisateur, m'inscrire et me connecter en sécurité (JWT).", "Haute"],
    ["US-02", "En tant qu'administrateur, gérer les utilisateurs, rôles et permissions fines.", "Haute"],
    ["US-03", "En tant que responsable, créer et configurer un programme et ses critères.", "Haute"],
    ["US-04", "En tant que responsable, construire un parcours visuel de sessions typées et imbriquées.", "Haute"],
    ["US-05", "En tant que responsable, définir le type et la visibilité de chaque session.", "Haute"],
    ["US-06", "En tant que responsable, voir le statut du programme se synchroniser depuis ses sessions.", "Moyenne"],
], widths=[1.3, 12.7, 2.0], cap_text="Backlog de la Release 1", font_size=9.5)
H2("4.2 Analyse")
P("Quatre cas d'utilisation majeurs sont livrés par cette release. Ils sont décrits ci-dessous "
  "sous forme de descriptions textuelles.")
P("**UC-1 — S'authentifier**", space_after=2)
uc_table([
    ["Acteurs", "Tous les acteurs inscrits"],
    ["But", "Obtenir une session authentifiée (JWT) pour accéder aux fonctionnalités protégées par rôle."],
    ["Préconditions", "L'utilisateur possède un compte."],
    ["Postconditions", "Un JWT signé est émis et stocké ; l'utilisateur est redirigé vers son tableau de bord."],
    ["Scénario nominal", "1. L'utilisateur ouvre la page de connexion. 2. Il saisit son courriel et son mot de passe. 3. Le système valide les identifiants. 4. Le service d'authentification émet un JWT contenant ses rôles et permissions. 5. Le front-end stocke le jeton et charge l'espace de travail approprié."],
    ["Alternatives", "3a. Identifiants invalides → un message d'erreur est affiché."],
    ["Exceptions", "Le service d'authentification est indisponible → une erreur technique est affichée."],
], "UC-1 — S'authentifier")
P("**UC-3 — Concevoir un programme et ses critères**", space_after=2)
uc_table([
    ["Acteurs", "Administrateur, Responsable de programme"],
    ["But", "Créer et configurer un programme et ses critères d'évaluation."],
    ["Préconditions", "L'acteur est authentifié et autorisé (`programmes:update`)."],
    ["Postconditions", "Le programme et ses critères sont persistés."],
    ["Scénario nominal", "1. L'acteur crée un programme et renseigne ses informations (dates, secteurs, types d'organisation éligibles, chiffres clés). 2. Il définit le formulaire de candidature. 3. Il ajoute des critères d'évaluation pondérés. 4. Il enregistre."],
    ["Alternatives", "1a. Dates incohérentes → une erreur de validation est retournée."],
    ["Exceptions", "—"],
], "UC-3 — Concevoir un programme et ses critères")
P("**UC-4 — Construire le parcours visuel**", space_after=2)
uc_table([
    ["Acteurs", "Administrateur, Responsable de programme"],
    ["But", "Modéliser le programme comme une suite de sessions et d'activités."],
    ["Préconditions", "Le programme existe ; l'acteur est autorisé."],
    ["Postconditions", "Les sessions, leur imbrication et leurs activités sont persistées ; le statut du programme peut être recalculé."],
    ["Scénario nominal", "1. L'acteur ajoute une session (jour ou plage). 2. Il définit son type, sa visibilité, ses dates, son lieu et ses participants. 3. Il imbrique éventuellement une session-jour dans une plage. 4. Il ajoute un agenda d'activités à une session-jour. 5. La couche de validation vérifie les dates, les chevauchements et les contraintes d'activités. 6. L'acteur confirme la création et prévisualise les modifications sensibles."],
    ["Alternatives", "5a. Une date viole une règle → un code d'erreur explicite est retourné (ex. SESSION_OVERLAP_DETECTED). 5b. Déplacer une plage sortirait une session-jour imbriquée → les jours imbriqués sont déplacés avec la plage après confirmation."],
    ["Exceptions", "—"],
], "UC-4 — Construire le parcours visuel")
P("**UC-5 — Définir le type et la visibilité d'une session**", space_after=2)
uc_table([
    ["Acteurs", "Administrateur"],
    ["But", "Définir le type d'une session et qui peut la voir."],
    ["Préconditions", "La session existe."],
    ["Postconditions", "Le type et la visibilité sont persistés ; la visibilité est appliquée côté serveur."],
    ["Scénario nominal", "1. L'acteur ouvre la session. 2. Il choisit un type (candidature, présélection, pitch day, …). 3. Il choisit une visibilité (Visible / Interne / Privée). 4. Il enregistre ; le changement est prévisualisé au préalable."],
    ["Alternatives", "—"],
    ["Exceptions", "—"],
], "UC-5 — Définir le type et la visibilité d'une session")
H2("4.3 Conception")
P("Le modèle du domaine tourne principalement autour du **Programme** et de ses **sessions** — "
  "l'entité unifiée `ProgrammePhase` — présentées ci-dessous. L'entité de session est nommée "
  "`ProgrammePhase` pour des raisons historiques — la classe a conservé le nom d'un concept "
  "antérieur, plus simple — mais elle représente bien une session au sens du chapitre 3 : typée, "
  "contrôlée en accès, datable et imbriquable. Son auto-association porte cette imbrication : une "
  "session-jour appartient à une session-plage.")
figure('image9.png', "Diagramme de classes de conception : domaine central")
P("Les deux tableaux suivants détaillent les attributs des entités structurantes de cette release.")
table([
    ["Attribut", "Type", "Description"],
    ["id", "Long", "Identifiant unique"],
    ["title", "String", "Titre du programme"],
    ["description", "Text", "Description détaillée"],
    ["type", "enum", "PUBLIC / PRIVATE"],
    ["status", "enum", "DRAFT/OPEN/IN_PROGRESS/EVALUATION/CLOSED/CANCELLED/ARCHIVED"],
    ["startDate / endDate", "Date", "Fenêtre du programme"],
    ["applicationDeadline", "Date", "Dérivée de la date de fin de la session de candidature"],
    ["sectors", "List", "Secteurs ciblés"],
], widths=[3.6, 2.0, 10.4], cap_text="Dictionnaire de données — Programme", font_size=9.5)
table([
    ["Attribut", "Type", "Description"],
    ["id", "Long", "Identifiant unique"],
    ["title", "String", "Titre de la session"],
    ["sessionType", "enum", "CANDIDATURE_SUBMISSION … TRAINING_DAY"],
    ["visibility", "enum", "VISIBLE / HIDDEN / PRIVATE"],
    ["allowActivities", "Boolean", "Faux pour les sessions de candidature"],
    ["allowOverlap", "Boolean", "Autorise le chevauchement dans un même couloir"],
    ["durationKind", "String", "« day » ou « range »"],
    ["startDate / endDate", "Date", "Fenêtre de la session"],
    ["status", "enum", "UPCOMING / ACTIVE / COMPLETED"],
    ["lane", "String", "Couloir / piste"],
    ["parentSessionId", "Long", "Plage parente d'une session-jour imbriquée"],
], widths=[3.6, 2.0, 10.4], cap_text="Dictionnaire de données — ProgrammePhase (Session)", font_size=9.5)
P("Côté sécurité, l'authentification issue d'UC-1 repose sur un JWT sans état. Le diagramme de "
  "séquence ci-dessous en montre le déroulé : le service d'authentification émet un jeton signé "
  "portant rôles et permissions, que chaque service valide ensuite localement pour peupler son "
  "contexte de sécurité.")
figure('image15.jpeg', "Diagramme de séquence : authentification JWT sans état")
H2("4.4 Réalisation")
P("Cette section présente les principaux écrans livrés par la release. Le premier est la page de "
  "connexion, porte d'entrée commune aux deux front-ends.")
figure('image16.png', "Écran de connexion")
P("Le tableau de bord du programme fait office de hub central du back-office : il fait remonter "
  "les indicateurs clés (candidatures, évaluations, acceptés), les prochaines sessions à "
  "notifier, la liste des éléments manquants et les meilleurs candidats.")
figure('image17.png', "Tableau de bord du programme (back-office)")
P("Le constructeur de parcours, quant à lui, affiche les sessions sur une frise organisée par "
  "couloir — formes jour et plage, imbrication, et agenda horaire pour les sessions-jour.")
figure('image18.png', "Constructeur de parcours visuel (back-office)")
P("La vue détaillée d'une session s'ouvre désormais sur un onglet **Détails** bien structuré — "
  "type, dates, statut, visibilité, description, compteurs associés — tandis que les activités et "
  "les panneaux fonctionnels restent dans des onglets séparés, pour ne pas surcharger la page.")
figure('image19.png', "Détail d'une session (onglets Détails / Activités)")
H2("4.5 Revue et rétrospective")
P("**Revue de sprint.** L'incrément livré répond aux user stories US-01 à US-06 : "
  "authentification JWT opérationnelle, administration des rôles et permissions, création de "
  "programme et constructeur de parcours fonctionnels, avec synchronisation automatique du "
  "statut. La démonstration au Product Owner a validé le cœur métier de la plateforme.")
P("**Rétrospective.** Cette release a aussi révélé plusieurs difficultés, traitées au fil de l'eau :")
bullet("**Une bizarrerie de nommage tenace.** L'entité de session s'appelle `ProgrammePhase` et "
       "vit encore dans une table `programme_phases` — un nom hérité d'une version antérieure du "
       "projet, avant que les sessions ne deviennent l'objet riche, typé et imbriqué "
       "d'aujourd'hui. Lire « ProgrammePhase » dans le code et le traduire mentalement en "
       "« Session » à chaque fois a demandé plus d'habitude que prévu.")
bullet("**Évolution du modèle de session.** Le modèle était initialement conçu autour d'un type "
       "fixe par programme. Les retours ont vite montré qu'un même programme devait contenir des "
       "sessions de natures différentes, imbriquées, chacune avec sa visibilité. Le modèle a donc "
       "été généralisé en l'entité unifiée `ProgrammePhase`, extensible sans toucher au schéma.")
bullet("**Des rôles non-admin aux mauvais accès.** Au début, les contrôles de permission étaient "
       "trop grossiers : un compte jury ou porteur pouvait parfois atteindre un endpoint interdit, "
       "ou se trouver bloqué sur un endpoint légitime. Les corriger a imposé de reprendre les "
       "contrôleurs un à un et de remplacer les vérifications de rôle larges par des permissions "
       "fines par module (`hasAuthority(\"module:action\")`).")
bullet("**Un blocage lié aux dates sur le changement de type.** Changer simplement le **type** "
       "d'une session — sans aucune modification de date — déclenchait à tort une revalidation "
       "complète des contraintes de dates et de chevauchement. Correction : ne revalider les "
       "dates que lorsque la modification les touche réellement.")
bullet("**Chevauchement des sessions imbriquées.** Déplacer une plage contenant des jours "
       "imbriqués échouait dès que ces jours sortaient de la nouvelle fenêtre. Désormais, les "
       "jours imbriqués sont déplacés avec la plage — ou ramenés à ses bornes — après confirmation.")
bullet("**Latence perçue sur la page des tâches.** Le chargement initial attendait, par erreur, "
       "des données non essentielles à l'affichage de la liste. Elles sont maintenant chargées "
       "paresseusement, à la demande.")
H2("Conclusion")
P("La Release 1 livre les fondations : sécurité, comptes et rôles d'un côté, modèle de programme "
  "et parcours visuel de l'autre. Les difficultés rencontrées ont, au passage, solidifié le "
  "modèle de données et le contrôle d'accès. La release suivante s'appuie sur ces fondations pour "
  "livrer le flux de candidature, l'évaluation et — surtout — l'analyse IA du pitch.")
pagebreak()

# =========================================================================
# CHAPITRE 5 — RELEASE 2
# =========================================================================
H1("Chapitre 5 — Release 2 : Candidatures, évaluation et analyse IA du pitch")
H2("Introduction")
P("La deuxième release est la plus dense du projet. Elle couvre les sprints 3 à 5 et livre le "
  "flux complet de sélection : le **dépôt de candidature** en ligne, l'**affectation des jurys**, "
  "l'**évaluation par critères pondérés** — assistée par l'IA **Medi** — et, point culminant, "
  "l'**analyse automatique du pitch** : transcription verbatim, mesures acoustiques, description "
  "visuelle et notation calibrée. C'est ici que se concrétisent les fondements théoriques du "
  "chapitre 2.")
H2("5.1 Backlog de la release")
table([
    ["#", "User story", "Prio."],
    ["US-07", "En tant que porteur, postuler en ligne via le formulaire propre au programme.", "Haute"],
    ["US-08", "En tant que porteur, lier ou créer mon organisation et mon équipe.", "Moyenne"],
    ["US-09", "En tant qu'administrateur, affecter des jurys à une session d'évaluation.", "Haute"],
    ["US-10", "En tant que jury, noter une candidature par critères pondérés, par session.", "Haute"],
    ["US-11", "En tant que jury, obtenir un score IA (Medi) par critère, réutilisable.", "Haute"],
    ["US-12", "En tant que porteur, téléverser des pitchs (entraînement et final) jusqu'à 2 Go.", "Haute"],
    ["US-13", "En tant que porteur, recevoir une analyse mesurée et horodatée de mon pitch.", "Haute"],
], widths=[1.3, 12.7, 2.0], cap_text="Backlog de la Release 2", font_size=9.5)
H2("5.2 Analyse")
P("**UC-2 — Déposer une candidature**", space_after=2)
uc_table([
    ["Acteurs", "Porteur"],
    ["But", "Postuler à un programme ouvert."],
    ["Préconditions", "Le programme est ouvert (sa session de candidature est active et la date limite n'est pas passée) ; le porteur est authentifié."],
    ["Postconditions", "Une `Candidature` est créée au statut *PENDING* et liée à l'organisation du porteur."],
    ["Scénario nominal", "1. Le porteur ouvre la page du programme. 2. Il démarre une candidature et remplit le formulaire propre au programme. 3. Il lie ou crée son organisation et son équipe. 4. Il soumet. 5. Le système vérifie l'éligibilité et la date limite, crée la candidature et confirme."],
    ["Alternatives", "5a. La date limite est passée → le dépôt est refusé. 5b. Des champs requis manquent → des erreurs de validation sont retournées."],
    ["Exceptions", "Le service de candidature est indisponible → le dépôt n'est pas enregistré et l'utilisateur est prévenu."],
], "UC-2 — Déposer une candidature")
P("**UC-6 — Affecter un jury**", space_after=2)
uc_table([
    ["Acteurs", "Administrateur, Responsable de programme"],
    ["But", "Désigner les jurys d'une session d'évaluation."],
    ["Préconditions", "Une session d'évaluation existe ; les personnes à affecter sont connues."],
    ["Postconditions", "Les affectations de jury sont créées (avec un lien d'évaluation tokenisé)."],
    ["Scénario nominal", "1. L'acteur ouvre la session d'évaluation. 2. Il sélectionne des jurys depuis l'annuaire ou par courriel. 3. Il confirme l'affectation. 4. Le système enregistre les affectations et peut notifier les jurys."],
    ["Alternatives", "—"],
    ["Exceptions", "—"],
], "UC-6 — Affecter un jury")
P("**UC-7 — Évaluer une candidature (par session)**", space_after=2)
uc_table([
    ["Acteurs", "Jury, Administrateur"],
    ["But", "Noter une candidature pour une session d'évaluation donnée."],
    ["Préconditions", "Le jury est affecté à la session."],
    ["Postconditions", "Une `Evaluation` clé (candidature, jury, session) est stockée ; le score consolidé est mis à jour."],
    ["Scénario nominal", "1. Le jury ouvre une candidature affectée. 2. Il note chaque critère pondéré et ajoute un commentaire. 3. Il peut lancer **Medi** et réutiliser ses notes. 4. Il soumet."],
    ["Alternatives", "3a. Le jury est en désaccord avec l'IA et ajuste les notes manuellement."],
    ["Exceptions", "—"],
], "UC-7 — Évaluer une candidature (par session)")
P("**UC-8 — Noter une candidature par IA**", space_after=2)
uc_table([
    ["Acteurs", "Jury, Administrateur (acteur système : Medi)"],
    ["But", "Obtenir un score proposé par l'IA pour une candidature."],
    ["Préconditions", "La candidature existe ; l'appelant est autorisé."],
    ["Postconditions", "Le score IA est affiché (non persisté tant qu'il n'est pas réutilisé/soumis)."],
    ["Scénario nominal", "1. L'acteur demande un score IA. 2. Le service IA récupère la candidature, les critères et le contexte organisation/équipe. 3. Il construit un prompt et appelle le LLM. 4. Il parse le résultat JSON (notes par critère, score pondéré, recommandation, synthèse). 5. Il retourne le résultat à l'interface."],
    ["Alternatives", "4a. La réponse du LLM ne peut être parsée → un repli sûr est retourné."],
    ["Exceptions", "Le LLM est indisponible → une erreur est affichée et la notation manuelle reste possible."],
], "UC-8 — Noter une candidature par IA")
P("Les deux diagrammes de séquence système ci-dessous précisent, respectivement, le dépôt d'une "
  "candidature et l'évaluation assistée par IA.")
figure('image5.jpeg', "Diagramme de séquence système : dépôt de candidature (UC-2)")
figure('image6.jpeg', "Diagramme de séquence système : évaluation assistée par IA (UC-7, UC-8)")
P("Le diagramme d'activité ci-dessous replace ces cas dans le flux global de sélection, du dépôt "
  "à l'intégration de la startup retenue.")
figure('image7.png', "Diagramme d'activité : candidature → sélection → intégration")
H2("5.3 Conception")
P("Deux entités structurent cette release : la **Candidature** et son **Evaluation** par session.")
table([
    ["Attribut", "Type", "Description"],
    ["id", "Long", "Identifiant unique"],
    ["programmeId", "Long", "Programme cible"],
    ["organizationId", "Long", "Organisation candidate"],
    ["carrierEmail", "String", "Courriel du porteur"],
    ["status", "enum", "PENDING … ACCEPTED/REJECTED"],
    ["totalScore", "Double", "Score pondéré consolidé"],
    ["customAnswers", "JSON", "Réponses au formulaire propre au programme"],
], widths=[3.6, 2.0, 10.4], cap_text="Dictionnaire de données — Candidature", font_size=9.5)
table([
    ["Attribut", "Type", "Description"],
    ["id", "Long", "Identifiant unique"],
    ["juryId", "Long", "Jury évaluateur"],
    ["sessionId", "Long", "Session d'évaluation"],
    ["weightedScore", "Double", "0–10, pondéré par les critères"],
    ["comment", "Text", "Commentaire libre"],
], widths=[3.6, 2.0, 10.4], cap_text="Dictionnaire de données — Evaluation", font_size=9.5)
P("L'analyse IA du pitch introduit deux structures dédiées. La première isole la **soumission de "
  "pitch**, qui porte la vidéo, sa transcription mise en cache et le résultat de l'analyse.")
figure('image10.png', "Diagramme de classes de conception : soumission de pitch")
P("Une soumission est liée à une session marquée pour analyse vidéo et typée par sa nature : un "
  "porteur peut déposer plusieurs vidéos TRAINING, pour mesurer la progression d'une tentative à "
  "l'autre, mais un seul pitch FINAL. La transcription et ses segments sont stockés sur la "
  "soumission, de sorte que relancer une analyse ne coûte que l'appel au modèle de langage, "
  "l'étape coûteuse de reconnaissance vocale étant réutilisée.")
P("La seconde structure détaille le **résultat de l'analyse**, dont la richesse distingue cette "
  "plateforme d'une simple note.")
figure('image11.png', "Diagramme de classes de conception : résultat d'analyse IA")
P("La figure ci-dessus rend visible dans le modèle le principe d'ancrage de la section 2.4.2. "
  "`DeliveryMetrics` est stéréotypée « **mesurée** » : chacun de ses champs provient du signal — "
  "ffmpeg pour l'intensité et les silences, l'alignement au niveau du mot pour les hésitations — "
  "et aucun n'est produit par le modèle de langage, dont le rôle se limite à les interpréter. Le "
  "score global n'est pas non plus un champ que le modèle remplit : il est dérivé côté serveur "
  "comme la moyenne des scores de `Dimension`, ce qui garde la note réconciliable avec le profil "
  "montré à l'utilisateur.")
H2("5.4 Réalisation")
P("Le tableau de bord d'évaluation réunit, pour l'administrateur et le jury, la revue d'une "
  "candidature et sa notation, avec le score IA de Medi affiché en regard des notes humaines.")
figure('image20.png', "Tableau de bord d'évaluation et revue de candidature (avec score IA)")
P("Côté front-office, le porteur remplit le formulaire de candidature propre au programme, "
  "tandis que le jury dispose d'une page d'évaluation dédiée.")
figure('image23.png', "Formulaire de candidature (front-office)")
figure('image24.png', "Page d'évaluation du jury (front-office)")
H2("5.5 Revue et rétrospective")
P("**Revue de sprint.** L'incrément livre le flux de sélection de bout en bout (US-07 à US-13) : "
  "candidature, affectation de jury, évaluation par session, score IA Medi et — surtout — "
  "l'analyse complète du pitch. C'est la release qui a apporté le plus de valeur différenciante, "
  "et aussi celle qui a exigé le plus de corrections.")
P("**Rétrospective.** Les difficultés rencontrées, particulièrement instructives sur la partie IA :")
bullet("**Intégration du LLM.** Fiabiliser le score IA a exigé un mécanisme de repli — des "
       "valeurs par défaut dès que la réponse du modèle n'était pas bien formée — pour que le "
       "jury ne soit jamais bloqué.")
bullet("**Une incompatibilité de protocole entre deux services.** Le service Java appelait le "
       "service média Python et recevait un 422 opaque. La cause n'avait rien à voir avec la "
       "charge utile : le client HTTP de Java négocie HTTP/2 par défaut, alors que le serveur "
       "Python ne parle que HTTP/1.1. Fixer la version du protocole a réglé le problème.")
bullet("**Une limite de téléversement qui ignorait la configuration.** Des vidéos bien en deçà "
       "de la limite étaient rejetées avec « Maximum upload size exceeded ». Le fichier de "
       "configuration était correct ; les variables d'environnement de docker-compose l'écrasaient "
       "silencieusement. La limite a été portée à 2 Go de façon cohérente sur la passerelle, le "
       "service et le front-end.")
bullet("**Un modèle de transcription qui supprime les hésitations.** Whisper est entraîné à "
       "produire une prose propre et retire chaque « euh » et « hum », ce qui rendait tout l'enjeu "
       "de l'analyse impossible. La première hypothèse — le filtre d'activité vocale — a été "
       "testée et invalidée : les deux réglages donnaient le même résultat. Amorcer le modèle avec "
       "un échantillon de contexte délibérément disfluent le fait transcrire verbatim, récupérant "
       "59 hésitations sur le même audio tout en préservant le contenu.")
bullet("**Des mesures qui paraissaient justes et ne l'étaient pas.** Le détecteur signalait 70 "
       "sons étirés, mais les mots concernés étaient « le », « et », « à » — des mots-outils. "
       "Whisper impute à un mot le silence qui le suit, si bien qu'un mot avant une pause paraît "
       "étiré. Soustraire les silences mesurés sur l'audio brut en a laissé 17 authentiques. Un "
       "nombre affiché avec l'autorité d'une mesure objective est pire que pas de nombre du tout.")
bullet("**Un modèle de vision non reproductible.** Le modèle rapportait « bras croisés » pour un "
       "orateur dont la main était levée — vérifié en extrayant la trame. La même trame, exécutée "
       "cinq fois, produisait des réponses différentes : le modèle échantillonnait. Sa température "
       "et sa graine ont été fixées, et les observations sur une seule trame sont désormais "
       "traitées comme des indices faibles.")
bullet("**Une échelle de notation qui s'est effondrée.** Averti que la notation était trop "
       "indulgente, la grille a été durcie avec des plafonds absolus (« si l'équipe ou la demande "
       "manque, la note ne peut dépasser 4 »). La correction est allée trop loin : un exposé "
       "délibérément mauvais et un vainqueur de concours de pitch obtenaient tous deux 3/10, si "
       "bien que l'échelle ne discriminait plus. Les plafonds ont été remplacés par des "
       "déductions par dimension, les attentes calibrées sur le format de pitch détecté, et le "
       "score global calculé côté serveur comme la moyenne des dimensions. Sur cinq vraies "
       "vidéos, l'écart est passé de 1,4 à 5 points.")
bullet("**Un fournisseur de modèle externe comme mode de défaillance.** Les boutons de génération "
       "se figeaient indéfiniment. Le code n'était pas en cause : le modèle configuré avait cessé "
       "de répondre. Deux défauts faisaient passer une panne transitoire pour permanente — "
       "l'absence de délai d'expiration côté client, et un délai serveur de quatre minutes hérité "
       "d'un tout autre usage. Les petites générations échouent désormais rapidement, avec un "
       "message clair.")
H2("Conclusion")
P("La Release 2 livre le cœur de la valeur du produit : la sélection assistée et, surtout, une "
  "analyse de pitch dont chaque chiffre est issu d'une mesure et non d'une supposition. Les "
  "corrections apportées sur la partie IA — transcription verbatim, silences soustraits, vision "
  "rendue reproductible, notation recalibrée — constituent le socle expérimental que le chapitre "
  "7 mettra à l'épreuve. La release suivante clôt le périmètre fonctionnel avec la visibilité, "
  "les participants et les communications.")
pagebreak()

# =========================================================================
# CHAPITRE 6 — RELEASE 3
# =========================================================================
H1("Chapitre 6 — Release 3 : Visibilité, participants et communications")
H2("Introduction")
P("La troisième release complète le périmètre fonctionnel. Elle couvre les sprints 6 et 7 : "
  "l'**application de la visibilité** côté serveur, adossée à la couche de validation "
  "centralisée, puis la **gestion des participants** et les **communications tracées** — "
  "invitations personnalisées, notifications par session et archive complète des envois. À "
  "l'issue de cette release, un programme peut être privé, ses participants gérés et informés, et "
  "chaque message rendu auditable.")
H2("6.1 Backlog de la release")
table([
    ["#", "User story", "Prio."],
    ["US-14", "En tant qu'administrateur, gérer les participants et envoyer des invitations tracées.", "Haute"],
    ["US-15", "En tant qu'administrateur, notifier les participants d'une session et archiver les envois.", "Moyenne"],
    ["US-16", "En tant que porteur / jury, suivre le parcours, mes tâches et mes mises à jour.", "Moyenne"],
    ["US-17", "En tant que visiteur, parcourir les programmes publics et leurs sessions visibles.", "Basse"],
    ["US-18", "En tant qu'administrateur, gérer les catalogues de référence.", "Basse"],
], widths=[1.3, 12.7, 2.0], cap_text="Backlog de la Release 3", font_size=9.5)
H2("6.2 Analyse")
P("**UC-9 — Gérer les participants et envoyer des invitations**", space_after=2)
uc_table([
    ["Acteurs", "Administrateur"],
    ["But", "Consulter les personnes associées à un programme et les inviter par un courriel tracé."],
    ["Préconditions", "L'acteur est autorisé."],
    ["Postconditions", "Les invitations sont envoyées et archivées (statut SENT/FAILED)."],
    ["Scénario nominal", "1. L'acteur ouvre la vue Participants. 2. Il parcourt l'annuaire (organisations, membres, porteurs, jurys). 3. Il ajoute des utilisateurs inscrits ou saisit des courriels. 4. Il personnalise le modèle d'invitation (variables, lien de suivi) et le prévisualise. 5. Il envoie ; chaque envoi est archivé."],
    ["Alternatives", "5a. Un courriel échoue → il est archivé au statut FAILED."],
    ["Exceptions", "—"],
], "UC-9 — Gérer les participants et envoyer des invitations")
P("**UC-10 — Suivre un programme**", space_after=2)
uc_table([
    ["Acteurs", "Porteur, Jury"],
    ["But", "Suivre le parcours d'un programme, ses mises à jour et ses tâches personnelles."],
    ["Préconditions", "L'acteur est authentifié et lié au programme."],
    ["Postconditions", "—"],
    ["Scénario nominal", "1. L'acteur ouvre le programme. 2. Il voit le parcours (sessions Visibles et, pour les utilisateurs invités, les sessions auxquelles ils sont conviés). 3. Il consulte ses tâches et ses mises à jour."],
    ["Alternatives", "—"],
    ["Exceptions", "—"],
], "UC-10 — Suivre un programme")
P("Le diagramme d'activité ci-dessous décrit la notification des participants d'une session, du "
  "choix des destinataires à l'archivage de chaque envoi.")
figure('image8.png', "Diagramme d'activité : notifier les participants d'une session")
H2("6.3 Conception")
P("L'entité **Invitation** est au centre de cette release : elle sert aussi bien aux invitations "
  "qu'aux notifications de session tracées, et porte un jeton de suivi unique.")
table([
    ["Attribut", "Type", "Description"],
    ["id", "Long", "Identifiant unique"],
    ["type", "enum", "JURY / PORTEUR / MEMBER / ORGANISATEUR / GUEST / GENERAL"],
    ["status", "enum", "PENDING / SENT / FAILED / ACCEPTED / DECLINED"],
    ["recipientEmail", "String", "Destinataire"],
    ["subject", "String", "Objet du courriel"],
    ["token", "String", "Jeton de suivi unique"],
    ["sentAt", "DateTime", "Horodatage d'envoi"],
], widths=[3.6, 2.0, 10.4], cap_text="Dictionnaire de données — Invitation", font_size=9.5)
P("Le diagramme de séquence ci-dessous détaille la notification tracée d'une session : "
  "récupération des destinataires, prévisualisation du modèle par type, envoi et archivage de "
  "chaque courriel avec son statut.")
figure('image12.jpeg', "Diagramme de séquence de conception : notification de session tracée")
P("Le dernier diagramme mérite qu'on s'y attarde, car c'est lui qui applique réellement la "
  "visibilité côté serveur — un point sur lequel repose une bonne part de la sécurité de la "
  "plateforme. Selon le rôle du demandeur, le service renvoie soit toutes les sessions "
  "(administrateur, responsable), soit uniquement les sessions Visibles augmentées de celles "
  "auxquelles le demandeur est invité.")
figure('image13.jpeg', "Diagramme de séquence de conception : filtrage des sessions par visibilité")
H2("6.4 Réalisation")
P("Le hub des participants réunit l'annuaire du programme et l'éditeur de modèle d'invitation, "
  "avec ses variables et son lien de suivi ; l'aperçu permet de vérifier le rendu avant l'envoi.")
figure('image21.png', "Hub des participants et modèle d'invitation (variables + lien de suivi)")
P("Côté front-office, la page publique du programme réunit les informations publiques, les "
  "sessions — avec leurs **badges de type**, toujours dans le respect de la **visibilité** — "
  "ainsi que les critères et les partenaires.")
figure('image22.png', "Page publique du programme (front-office)")
H2("6.5 Revue et rétrospective")
P("**Revue de sprint.** L'incrément livre US-14 à US-18 : gestion des participants, invitations "
  "et notifications tracées avec archive, suivi de programme côté porteur et jury, consultation "
  "publique respectant la visibilité, et catalogues de référence. Le périmètre fonctionnel "
  "initial est désormais complet.")
P("**Rétrospective.** Cette release, plus courte, a surtout consolidé l'existant. Le principal "
  "enseignement porte sur la **couche de validation centralisée** : en réunissant en un seul "
  "endroit les règles de dates, de chevauchement et d'imbrication, avec des codes d'erreur "
  "explicites, elle a rendu le comportement du système prévisible et testable — condition "
  "nécessaire aux scénarios de validation du chapitre suivant. L'application de la visibilité "
  "côté serveur, plutôt que dans la seule interface, a par ailleurs été confirmée comme la "
  "garantie sur laquelle repose la fonctionnalité de programme privé.")
H2("Conclusion")
P("Avec la Release 3, la plateforme couvre l'intégralité du cycle de vie d'un programme "
  "d'incubation : conception, candidatures, évaluation, analyse de pitch, visibilité, "
  "participants et communications. Il reste à démontrer, mesures à l'appui, que l'ensemble tient "
  "ses promesses. C'est l'objet du chapitre suivant.")
pagebreak()

# =========================================================================
# CHAPITRE 7 — VALIDATION, TESTS ET DÉPLOIEMENT
# =========================================================================
H1("Chapitre 7 — Validation, tests et déploiement")
H2("Introduction")
P("Les chapitres précédents ont décrit ce qui a été construit ; celui-ci se demande si cela "
  "fonctionne, et répond par des mesures plutôt que par des affirmations. La question importe "
  "plus ici que dans une application ordinaire : une plateforme qui note le pitch d'un fondateur "
  "inspire confiance, et un outil auquel on fait confiance tout en ayant tort est pire que pas "
  "d'outil du tout. Ce chapitre, qui correspond au Sprint 8, expose donc le protocole de "
  "validation, les résultats obtenus sur la fiabilité des mesures, la calibration de la notation "
  "et la performance, puis les scénarios de test fonctionnels et la planification globale du projet.")
H2("7.1 Protocole de validation")
P("Valider un outil d'évaluation ne peut se faire sur un seul enregistrement : une note n'a de "
  "sens que relativement à d'autres notes. Un corpus de référence de cinq vraies vidéos a donc "
  "été constitué, couvrant délibérément tout le spectre de qualité, de sorte que le classement "
  "attendu soit connu à l'avance et indépendamment de la plateforme.")
table([
    ["Vidéo de référence", "Niveau attendu", "Objectif"],
    ["Pitch YC Demo Day (Goodybag)", "Excellent", "Ancre haute : forte traction avérée"],
    ["Vainqueur de concours de pitch", "Bon", "Doit se classer nettement au-dessus d'un mauvais exposé"],
    ["Vainqueur d'elevator pitch", "Bon", "Format très court (57 s)"],
    ["Exposé étudiant délibérément mauvais", "Faible", "Ancre basse"],
    ["Parodie (« How to sound smart »)", "Pas un pitch", "Détecte le cas dégénéré"],
], widths=[4.6, 3.0, 8.4], cap_text="Corpus vidéo de référence utilisé pour la validation", font_size=9.5)
P("Le critère de succès n'est pas une note absolue — il n'existe pas de vérité terrain pour « ce "
  "pitch mérite 7/10 » — mais la **discrimination** : le classement produit doit correspondre au "
  "classement connu, et l'écart entre le meilleur et le pire doit être assez large pour qu'un "
  "jury puisse agir. Ce corpus est rejoué après chaque changement de la grille, ce qui en fait "
  "un test de non-régression plutôt qu'une expérience ponctuelle.")
H2("7.2 Fiabilité des mesures")
P("Puisque toute l'évaluation est ancrée dans des mesures, les mesures elles-mêmes ont été "
  "validées en premier. Trois résultats sont rapportés, chacun correspondant à un défaut trouvé "
  "et corrigé.")
bullet("**Récupérer les hésitations.** Comme prévu en section 2.1.2, la transcription ne "
       "contenait initialement aucune pause remplie : zéro sur un enregistrement qui en contient "
       "audiblement beaucoup. L'hypothèse du filtre d'activité vocale a été testée et rejetée — "
       "les deux réglages produisaient une sortie strictement identique. Amorcer le décodeur avec "
       "un échantillon de contexte disfluent a récupéré 59 hésitations sur le même audio, tout en "
       "préservant le contenu (0,94 de similarité mot à mot avec la transcription propre). La "
       "technique ne coûte aucun temps supplémentaire.")
bullet("**Éliminer les faux positifs.** Le premier détecteur de prolongations signalait 70 sons "
       "étirés. L'inspection du détail montrait que les mots concernés étaient « le », « et », "
       "« à » — des mots-outils. La cause est la propriété d'alignement décrite en section 2.2.3 : "
       "le silence qui suit un mot lui est imputé. Soustraire les silences mesurés sur l'audio "
       "brut a laissé 17 prolongations authentiques, une réduction de 76 %.")
bullet("**Reproductibilité de l'observation visuelle.** Le modèle de vision décrivait un orateur "
       "comme ayant les bras croisés à 2:17 ; extraire la trame montrait sa main levée. Exécuté "
       "cinq fois sur cette trame identique, le modèle produisait des descriptions différentes. "
       "Fixer la température et la graine aléatoire a rendu la sortie reproductible et correcte "
       "sur trois exécutions consécutives.")
table([
    ["Mesure", "Avant correction", "Après correction"],
    ["Hésitations détectées (même audio)", "0", "59"],
    ["Prolongations détectées", "70 (mots-outils)", "17 (authentiques)"],
    ["Vision : même trame, 5 exécutions", "incohérent", "identique et correct"],
], widths=[6.0, 5.0, 5.0], cap_text="Effet des corrections sur la fiabilité des mesures", font_size=9.5)
H2("7.3 Calibration de l'échelle de notation")
P("La première grille a été jugée trop indulgente. La correction appliquée — des plafonds "
  "absolus du type « si l'équipe ou la demande de financement manque, la note ne peut dépasser "
  "4 » — a produit une défaillance pire, que le corpus de référence a révélée immédiatement : un "
  "exposé délibérément mauvais et un vainqueur de concours de pitch obtenaient tous deux 3/10. "
  "Les plafonds s'appliquaient avant que les points forts ne soient lus. Une échelle qui ne peut "
  "séparer le meilleur du pire ne porte aucune information.")
P("Trois changements ont été faits : un manque abaisse désormais seulement la dimension "
  "concernée au lieu de plafonner le score global ; les attentes sont calibrées sur le format de "
  "pitch détecté ; et le score global est dérivé côté serveur comme la moyenne des dimensions. "
  "Le corpus a ensuite été rejoué.")
table([
    ["Vidéo de référence", "Format détecté", "Avant", "Après"],
    ["YC Demo Day (Goodybag)", "Demo Day", "4", "7"],
    ["Vainqueur d'elevator pitch", "Demo Day", "4", "5,9"],
    ["Vainqueur de concours", "Compétition", "3", "5"],
    ["Mauvais exposé étudiant", "Pas un pitch", "3", "2,3"],
    ["Parodie", "Pas un pitch", "2,6", "2"],
    ["Écart (max − min)", "—", "1,4", "5,0"],
], widths=[5.5, 4.5, 3.0, 3.0], cap_text="Calibration de la notation mesurée sur le corpus de référence", font_size=9.5)
P("Le résultat est sans ambiguïté : le pouvoir discriminant est passé de 1,4 à 5,0 points sur "
  "10, le classement correspond désormais à celui attendu, et les deux pitchs gagnants se "
  "classent enfin au-dessus du mauvais exposé et de la parodie. Rien dans les vidéos n'a changé — "
  "seulement la grille. **C'est le résultat expérimental central du projet**, et il valide le "
  "principe de calibration énoncé en section 2.4.2.")
H2("7.4 Performance")
P("La première version du pipeline traitait une vidéo en environ 89 secondes, en enchaînant "
  "transcription, analyse audio et vision. Ces trois traitements étant indépendants, ils ont été "
  "parallélisés, ramenant le temps de traitement à environ 13 secondes — un facteur d'environ "
  "sept. Mettre la transcription en cache supprime entièrement l'étape la plus coûteuse quand une "
  "analyse est rejouée.")
P("Un second résultat concerne le modèle de langage. Le fournisseur initialement configuré avait "
  "cessé de répondre et chaque appel expirait. Le benchmark des modèles disponibles sur la même "
  "clé a montré un modèle léger répondant en 0,5 s là où le grand expirait, ce qui a tranché le "
  "choix. Deux défauts ont été corrigés en même temps : l'absence de délai d'expiration côté "
  "client et un délai serveur de quatre minutes hérité d'un tout autre usage.")
table([
    ["Indicateur", "Avant", "Après"],
    ["Traitement vidéo (pitch de 5 min)", "≈ 89 s", "≈ 13 s"],
    ["Analyse rejouée (transcription en cache)", "pipeline complet", "appel LLM seul"],
    ["Génération de champ (assistance IA)", "expiration", "≈ 0,5 – 8 s"],
], widths=[6.0, 5.0, 5.0], cap_text="Performance avant et après optimisation", font_size=9.5)
H2("7.5 Jeu de données de test")
P("Pour valider la plateforme, un jeu de données représentatif a été constitué : au moins un "
  "compte par rôle ; plusieurs organisations avec leurs membres et porteurs ; un programme "
  "complet avec un parcours couvrant chaque type de session, les deux formes (jour/plage), "
  "l'imbrication et les trois niveaux de visibilité ; une cohorte de candidatures avec réponses "
  "personnalisées ; des affectations de jury et scores pondérés, plus des scores IA (Medi) pour "
  "comparaison ; et des invitations et notifications tracées avec entrées d'archive SENT/FAILED. "
  "Le tableau ci-dessous liste les scénarios de test métier exécutés sur ce jeu de données.")
table([
    ["#", "Scénario", "Résultat attendu"],
    ["T1", "Créer une plage dont un jour imbriqué sortirait des bornes", "Erreur explicite, ou le jour imbriqué est déplacé avec la plage"],
    ["T2", "Deux sessions qui se chevauchent dans le même couloir", "SESSION_OVERLAP_DETECTED (sauf chevauchement autorisé)"],
    ["T3", "Un utilisateur non privilégié liste les sessions", "Seules les sessions VISIBLE (et invitées) sont retournées"],
    ["T4", "Déposer une candidature après la date limite", "Dépôt refusé"],
    ["T5", "Un jury note le même candidat dans deux sessions d'évaluation", "Deux évaluations indépendantes sont enregistrées"],
    ["T6", "Noter une candidature par IA", "Notes par critère + score pondéré + recommandation"],
    ["T7", "Envoyer des notifications de session", "Courriels par type envoyés + archivés (SENT/FAILED)"],
    ["T8", "Changer le type d'une session", "Changement persisté et reflété comme badge sur le front-office"],
], widths=[1.0, 7.0, 8.0], cap_text="Scénarios de test métier", font_size=9.5)
H2("7.6 Scénarios de test fonctionnels")
P("En parallèle de la validation du pipeline, les principaux chemins fonctionnels ont été "
  "vérifiés de bout en bout sur la plateforme déployée, chacun via la passerelle d'API avec un "
  "compte et un jeton réels.")
table([
    ["Scénario", "Résultat attendu", "Statut"],
    ["Connexion et émission du JWT", "Jeton portant rôles et permissions", "Réussi"],
    ["Changement de droits en direct (SSE)", "Interface rafraîchie sans reconnexion", "Réussi"],
    ["Création d'un programme en brouillon, reprise plus tard", "Chaque étape persistée et restaurée", "Réussi"],
    ["La session de candidature pilote la date de clôture", "Date limite dérivée automatiquement", "Réussi"],
    ["Téléversement vidéo (jusqu'à 2 Go)", "Stockée et liée à la soumission", "Réussi"],
    ["Analyse complète du pitch", "Score, moments-clés et coaching retournés", "Réussi"],
    ["Visibilité d'un programme privé", "Listé uniquement pour les utilisateurs invités", "Réussi"],
], widths=[6.5, 6.5, 3.0], cap_text="Principaux scénarios de test fonctionnels", font_size=9.5)
H2("7.7 Planification et déploiement")
P("Le développement s'est déployé, de façon reproductible, via **Docker Compose** : chaque "
  "service et chaque base est un conteneur, l'ensemble étant démarrable d'une seule commande. "
  "Cette conteneurisation, posée dès le Sprint 0, a garanti que l'environnement d'exécution reste "
  "identique du poste de développement à l'intégration. La feuille de route Scrum ci-dessous "
  "retrace la planification réelle du projet : trois releases regroupant neuf sprints, chacun "
  "livrant un incrément déployable.")
figure(os.path.join(HERE, 'planning_scrum.png'),
       "Feuille de route Scrum du projet (releases, sprints et incréments)")
H2("Conclusion")
P("Ce chapitre a répondu à la question qu'il posait en ouverture. Les mesures ont été validées "
  "une à une, et chaque validation a intercepté un défaut qui aurait autrement atteint "
  "l'utilisateur : un compte d'hésitations absent, un compte de prolongations gonflé, une posture "
  "inventée. L'échelle de notation, testée contre un corpus dont le classement était connu "
  "d'avance, est passée d'un état où elle ne séparait rien à un écart de 5 points sur 10. La "
  "performance a été divisée par sept. Le résultat essentiel est méthodologique : chaque "
  "affirmation de la plateforme sur un pitch est désormais traçable jusqu'à une mesure, et cette "
  "mesure a elle-même été vérifiée contre sa source.")
pagebreak()

# =========================================================================
# CONCLUSION GÉNÉRALE
# =========================================================================
H1("Conclusion générale")
P("**Récapitulation de la démarche.** L'introduction a posé le problème d'emblée : les "
  "programmes d'incubation souffrent, en pratique, d'une réelle fragmentation. La démarche s'est "
  "déroulée de façon agile. Après avoir posé le cadre et étudié l'existant (chapitre 1) et "
  "établi les fondements théoriques de l'analyse de pitch (chapitre 2), le Sprint 0 a produit un "
  "backlog produit priorisé et une architecture microservices (chapitre 3). Trois releases "
  "successives ont ensuite livré, par incréments validés, les fondations et le parcours "
  "(chapitre 4), la sélection et l'analyse IA du pitch (chapitre 5), puis la visibilité, les "
  "participants et les communications (chapitre 6). Le tout a enfin été validé, mesures à "
  "l'appui (chapitre 7).")
P("**Présentation des résultats.** Les résultats répondent-ils au problème posé au départ ? "
  "Largement, oui. **Medianet Incubateur** réunit désormais, dans un outil sécurisé unique, la "
  "conception visuelle d'un programme, les candidatures en ligne, l'évaluation par jury assistée "
  "par IA, la gestion des participants et des organisations, et des communications réellement "
  "traçables. La visibilité des sessions, appliquée côté serveur comme dans l'interface, avec une "
  "couche de validation centralisée à codes d'erreur explicites, apporte la fiabilité et le "
  "contrôle d'accès qui manquaient aux solutions existantes.")
P("**Problèmes rencontrés lors de la réalisation.** Les sections de rétrospective les détaillent, "
  "de la bizarrerie de nommage `ProgrammePhase` aux permissions trop grossières, en passant par "
  "la généralisation tardive du modèle de session et les défauts de la partie IA. Aucun n'a exigé "
  "de repenser l'architecture — chaque cas a été une correction ciblée du modèle, du contrôle "
  "d'accès ou de la validation.")
P("**Apports.** Techniquement, ce projet a signifié parcourir de bout en bout l'ingénierie d'un "
  "système distribué non trivial : modélisation du domaine et conception UML, architecture "
  "microservices avec Spring Cloud, sécurité sans état avec contrôle d'accès par rôles, "
  "développement front-end moderne en Next.js et TypeScript, intégration d'un LLM externe, DevOps "
  "avec Docker Compose. L'apport sur lequel le projet se défend le plus, toutefois, est "
  "méthodologique plutôt que purement technique : il montre qu'**une évaluation par IA ne vaut "
  "que ce que valent ses mesures**. Un modèle de transcription qui effaçait silencieusement "
  "chaque hésitation rendait l'analyse dénuée de sens ; un détecteur signalant soixante-dix sons "
  "étirés comptait des mots ordinaires suivis d'une pause ; un modèle de vision décrivait avec "
  "assurance un orateur bras croisés quand la trame montrait sa main levée. Dans les trois cas, "
  "un nombre plausible et présenté avec assurance était simplement faux. Le travail de "
  "calibration qui a suivi a été validé sur un corpus de vraies vidéos : l'écart entre le "
  "meilleur et le pire pitch est passé de 1,4 à 5 points sur 10.")
P("**Perspectives.** Plusieurs pistes méritent d'être poursuivies. Les types de session "
  "pourraient devenir entièrement pilotés par catalogue — migrant d'une énumération figée vers "
  "une valeur référencée, pour que de futurs types (hackathons, journées investisseurs, …) "
  "n'exigent aucune modification de code. Les notifications temps réel (WebSocket / SSE) et "
  "l'export calendrier (`.ics` / Outlook) sont une autre direction naturelle, tout comme le "
  "renforcement de l'**observabilité** (journalisation centralisée, métriques, traçage) et la "
  "construction d'un vrai pipeline **CI/CD** avec tests de bout en bout. Enfin, le volet "
  "**analytique** a de la marge : entonnoirs de cohortes, accord inter-jurys, délai de décision, "
  "voire un catalogue public et multilingue de programmes. Rien de tout cela n'était strictement "
  "requis pour répondre au besoin initial ; mais c'est précisément ce qu'une fondation durable "
  "doit rendre possible.")
H2("Travail quotidien et rôle dans l'équipe")
P("Au jour le jour, le travail a suivi un rythme régulier : un point court avec l'encadrant sur "
  "la priorité du moment, puis le développement proprement dit — lire le code existant avant de "
  "le modifier, implémenter, tester localement contre la pile complète sous Docker, et déployer "
  "sur l'environnement d'intégration. Une part récurrente de chaque journée allait au diagnostic "
  "plutôt qu'à l'écriture de code : reproduire un défaut, formuler une hypothèse, la tester, puis "
  "seulement corriger. Les retours de l'équipe, souvent donnés comme une observation brute — « ça "
  "a pris trop de temps », « l'IA est trop généreuse » — étaient systématiquement transformés en "
  "problème mesurable avant tout changement. Le rôle tenu était celui d'un développeur full-stack "
  "avec une réelle autonomie sur le périmètre confié.")
H2("Compétences acquises")
P("Techniquement, le projet a consolidé le développement back-end (Java, Spring Boot, sécurité, "
  "conception d'API REST), le développement front-end (React, Next.js, TypeScript), la "
  "modélisation relationnelle et ses pièges de migration, le déploiement par conteneurs, et "
  "surtout un domaine nouveau : l'intelligence artificielle appliquée — reconnaissance de la "
  "parole, traitement du signal audio, modèles de vision et intégration de modèles de langage, y "
  "compris la conception et la calibration d'une grille d'évaluation. Méthodologiquement, la "
  "compétence qui a le plus pesé est la rigueur de la preuve : ne jamais présenter un chiffre "
  "sans l'avoir confronté à sa source. Elle s'est apprise à la dure, sur des mesures qui étaient "
  "plausibles et fausses.")
H2("Apports mutuels et atteinte des objectifs")
P("**Ce que le projet a apporté à l'entreprise.** Medianet dispose d'une plateforme "
  "opérationnelle couvrant le cycle d'incubation et, surtout, d'un module différenciant : une "
  "analyse automatique de pitch entièrement auto-hébergée, donc utilisable sur les vidéos "
  "confidentielles des fondateurs sans qu'aucun service tiers ne soit impliqué. Le corpus de "
  "validation et le protocole bâti autour restent utilisables par l'équipe pour requalifier la "
  "grille après tout changement futur.")
P("**Ce que l'entreprise a apporté au projet.** Medianet a fourni un vrai sujet plutôt qu'un "
  "exercice — avec de vrais utilisateurs, de vraies contraintes et des retours exigeants — "
  "l'infrastructure nécessaire à l'exécution de modèles locaux, et un environnement technique où "
  "les décisions d'architecture devaient être justifiées.")
P("**Les objectifs ont-ils été atteints ?** Les objectifs fonctionnels fixés au départ sont "
  "atteints : la plateforme modélise un programme comme un parcours de sessions typées, gère les "
  "candidatures et leur évaluation, et applique le contrôle d'accès côté serveur. L'objectif "
  "attaché à la seconde problématique — rendre l'évaluation d'un pitch objective, reproductible "
  "et utile — est atteint et mesuré : l'analyse est ancrée dans des mesures vérifiées, elle est "
  "reproductible, elle discrimine (5 points d'écart contre 1,4), et elle ne quitte jamais "
  "l'infrastructure. Deux objectifs restent partiellement ouverts : l'analyse visuelle de la "
  "présence reste un indice faible, limité par le modèle compact utilisé, et l'évaluation de "
  "l'utilité réelle du coaching exigerait un suivi des porteurs sur un programme complet, que la "
  "durée du projet n'a pas permis.")
pagebreak()

# =========================================================================
# BIBLIOGRAPHIE
# =========================================================================
H1("Bibliographie / Webographie")
biblio = [
    "[1] Spring Boot Reference Documentation, Spring Team / VMware, en ligne : https://docs.spring.io/spring-boot/documentation.html (consulté en 2026).",
    "[2] Spring Cloud Gateway Reference Documentation, Spring Team / VMware, en ligne : https://docs.spring.io/spring-cloud-gateway/reference/ (consulté en 2026).",
    "[3] Spring Cloud Netflix (Eureka) Reference Documentation, Spring Team / VMware, en ligne : https://docs.spring.io/spring-cloud-netflix/reference/ (consulté en 2026).",
    "[4] Next.js Documentation, Vercel Inc., en ligne : https://nextjs.org/docs (consulté en 2026).",
    "[5] JSON Web Token (JWT) — RFC 7519, M. Jones, J. Bradley, N. Sakimura, IETF, mai 2015, en ligne : https://www.rfc-editor.org/rfc/rfc7519 (consulté en 2026).",
    "[6] M. Fowler, J. Lewis, « Microservices: a definition of this new architectural term », martinfowler.com, 2014, en ligne : https://martinfowler.com/articles/microservices.html (consulté en 2026).",
    "[7] A. Radford et al., « Robust Speech Recognition via Large-Scale Weak Supervision » (Whisper), OpenAI, 2022.",
    "[8] EBU R 128, « Loudness normalisation and permitted maximum level of audio signals », European Broadcasting Union ; ITU-R BS.1770.",
    "[9] OWASP Foundation, « OWASP Top Ten », en ligne : https://owasp.org/www-project-top-ten/ (consulté en 2026).",
    "[10] PostgreSQL Documentation, PostgreSQL Global Development Group, en ligne : https://www.postgresql.org/docs/ (consulté en 2026).",
    "[11] MinIO Documentation, MinIO Inc., en ligne : https://min.io/docs/minio/linux/index.html (consulté en 2026).",
    "[12] Docker Compose Documentation, Docker Inc., en ligne : https://docs.docker.com/compose/ (consulté en 2026).",
]
for b in biblio:
    P(b, indent=False, space_after=6)
pagebreak()

# =========================================================================
# ANNEXES
# =========================================================================
H1("Annexe A — Cartographie des services et des ports")
table([
    ["Service", "Port", "Responsabilité"],
    ["api-gateway", "8080", "Point d'entrée, routage, transmission du JWT"],
    ["eureka-server", "8761", "Découverte de services"],
    ["auth-service", "8081", "Utilisateurs, rôles, organisations, JWT, fichiers"],
    ["candidature-service", "8083", "Candidatures, évaluations, affectations de jury"],
    ["programme-service", "8086", "Programmes, sessions, critères, tâches, validation, visibilité"],
    ["notification-service", "8087", "Invitations, notifications, archive des courriels"],
    ["admin-ai-service", "8088", "Notation IA (Medi), assistant d'administration"],
    ["pitch-media-service", "8089", "Pipeline vidéo : transcription (Whisper), mesures audio (ffmpeg), vision"],
    ["ollama", "11434", "Runtime de modèles local (modèle de vision pour la présence physique)"],
    ["nextjs-frontoffice", "3000", "Front-end public / porteur / jury"],
    ["nextjs-backoffice", "3001", "Front-end d'administration"],
    ["PostgreSQL (×5)", "5432 / 5434–5437", "Une base de données par service"],
    ["MinIO", "9000 / 9001", "Stockage objet (compatible S3)"],
], widths=[3.6, 2.6, 9.8], font_size=9.5)
P("Les anciens `ai-scoring-service` (8084) et `ai-matching-service` (8085) ont été remplacés par "
  "Medi / admin-ai-service.", space_after=6)

H1("Annexe B — Principaux endpoints REST")
table([
    ["Méthode", "Chemin", "Description"],
    ["POST", "/api/auth/login", "S'authentifier et obtenir un JWT"],
    ["GET", "/api/programmes", "Lister les programmes publics"],
    ["GET", "/api/programmes/{id}", "Détail d'un programme (phases filtrées par visibilité)"],
    ["GET", "/api/programmes/{id}/sessions", "Lister les sessions (filtrées par visibilité)"],
    ["POST", "/api/programmes/{id}/sessions", "Créer une session"],
    ["PUT", "/api/programmes/{id}/sessions/{sid}", "Modifier une session"],
    ["POST", "/api/candidatures", "Déposer une candidature"],
    ["GET", "/api/candidatures/programme/{id}", "Candidatures d'un programme"],
    ["POST", "/api/candidatures/{id}/assign-jury", "Affecter des jurys à une session"],
    ["POST", "/api/candidatures/{id}/evaluate", "Soumettre une évaluation de jury"],
    ["POST", "/api/admin-ai/score/{id}", "Noter une candidature par IA"],
    ["POST", "/api/notifications/email/session-notify", "Envoyer des notifications de session tracées"],
    ["GET", "/api/notifications/invitations/programme/{id}", "Archive des communications"],
    ["POST", "/api/pitch/submissions", "Créer/mettre à jour une soumission de pitch"],
    ["GET", "/api/pitch/submissions/mine", "Les soumissions de pitch du porteur"],
    ["GET", "/api/pitch/presentations/{programmeId}", "Sessions de présentation + soumissions par session"],
    ["PUT", "/api/pitch/submissions/{id}/analysis", "Persister l'analyse IA sur une soumission"],
    ["POST", "/api/admin-ai/pitch/analyze/stream", "Lancer l'analyse du pitch en diffusant les étapes (SSE)"],
    ["POST", "/api/admin-ai/field-suggest", "Générer ou améliorer un champ de programme"],
    ["GET", "/api/programmes/invited", "Programmes privés auxquels l'appelant est invité"],
], widths=[2.0, 7.0, 7.0], cap_text="Extrait des principaux endpoints REST", font_size=9.0)

H1("Annexe C — Glossaire")
table([
    ["Terme", "Définition"],
    ["Parcours", "La frise visuelle des sessions et activités d'un programme."],
    ["Session (Phase)", "Une étape typée d'un programme (candidature, présélection, pitch day, …)."],
    ["Jour / Plage", "Une session sur une journée / une session sur une plage de dates."],
    ["Porteur", "Un porteur de projet / fondateur de startup qui postule."],
    ["Jury", "Un évaluateur qui note les candidatures."],
    ["Visibilité", "Le fait qu'une session soit Visible, Masquée (interne) ou Privée."],
    ["JWT", "JSON Web Token, l'identifiant d'authentification sans état."],
    ["RBAC", "Role-Based Access Control (contrôle d'accès basé sur les rôles)."],
    ["LUFS / LRA", "Unités d'intensité sonore (EBU R 128) / plage d'intensité."],
    ["ASR", "Automatic Speech Recognition (reconnaissance automatique de la parole)."],
    ["Medi", "Le service d'IA fournissant la notation et l'analyse de pitch."],
], widths=[3.4, 12.6], font_size=9.5)
pagebreak()

# --- Quatrième de couverture ESPRIT ---
fullpage_image('image26.png')

_set_updatefields()
doc.save(OUT)
print("saved", OUT)
